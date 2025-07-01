import streamlit as st
import tempfile
import os
from io import BytesIO
import base64
import getpass

# PDF
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

# URL Scraping
try:
    import requests
    from bs4 import BeautifulSoup
except ImportError:
    requests = None
    BeautifulSoup = None

# YouTube Transcript
try:
    from youtube_transcript_api import YouTubeTranscriptApi
except ImportError:
    YouTubeTranscriptApi = None

# Image OCR & Captioning
try:
    from PIL import Image
    import pytesseract
except ImportError:
    Image = None
    pytesseract = None
try:
    from transformers import BlipProcessor, BlipForConditionalGeneration
except ImportError:
    BlipProcessor = None
    BlipForConditionalGeneration = None

# Audio Transcription & Analysis
try:
    import whisper
except ImportError:
    whisper = None
try:
    from pyAudioAnalysis import audioSegmentation as aS
except ImportError:
    aS = None

# HuggingFace Transformers
try:
    from transformers import pipeline
except ImportError:
    pipeline = None

# spaCy
try:
    import spacy
    nlp = spacy.load("en_core_web_sm")
except Exception:
    nlp = None

# TTS
try:
    from gtts import gTTS
except ImportError:
    gTTS = None
try:
    import pyttsx3
except ImportError:
    pyttsx3 = None

# Knowledge Graph
try:
    import networkx as nx
    from pyvis.network import Network
except ImportError:
    nx = None
    Network = None

# Plagiarism/Similarity
try:
    from sentence_transformers import SentenceTransformer, util
    from sklearn.metrics.pairwise import cosine_similarity
except ImportError:
    SentenceTransformer = None
    util = None
    cosine_similarity = None

# Set page config and apply glassmorphic, ultra-modern CSS
st.set_page_config(page_title="Obsidian Protocol v2.0", layout="wide")
st.markdown('''
    <style>
    @import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700&display=swap');
    body, .stApp {
        background: linear-gradient(120deg, #1a1f2b 0%, #232946 50%, #2e3a59 100%);
        min-height: 100vh;
        animation: glassbgmove 24s ease-in-out infinite alternate;
        transition: background 0.8s cubic-bezier(.4,2,.6,1);
        will-change: background, opacity, transform;
    }
    @keyframes glassbgmove {
        0% { background-position: 0% 50%; }
        100% { background-position: 100% 50%; }
    }
    .aurora-bg {
        position: fixed; top: 0; left: 0; width: 100vw; height: 100vh; z-index: 0;
        pointer-events: none;
        background: radial-gradient(circle at 20% 40%, #00c3ff44 0%, #0057b822 60%, transparent 100%),
                    radial-gradient(circle at 80% 60%, #ffd70044 0%, #ffb30022 60%, transparent 100%);
        animation: auroraMove 18s ease-in-out infinite alternate;
        opacity: 0.7;
    }
    @keyframes auroraMove {
        0% { background-position: 0% 0%, 100% 100%; }
        100% { background-position: 100% 100%, 0% 0%; }
    }
    .glass-sidebar {
        position: fixed;
        top: 0; left: 0; bottom: 0;
        width: 72px;
        background: rgba(34, 40, 60, 0.22);
        box-shadow: 4px 0 32px #0057B822;
        backdrop-filter: blur(24px) saturate(180%);
        -webkit-backdrop-filter: blur(24px) saturate(180%);
        border-right: 2px solid rgba(255,255,255,0.08);
        z-index: 10001;
        display: flex;
        flex-direction: column;
        align-items: center;
        padding: 1.5rem 0 1.2rem 0;
        gap: 1.2rem;
        transition: width 0.35s cubic-bezier(.4,2,.6,1), background 0.5s, box-shadow 0.3s cubic-bezier(.4,2,.6,1);
    }
    .glass-sidebar.expanded {
        width: 260px;
        align-items: flex-start;
        padding-left: 1.2rem;
        background: rgba(34, 40, 60, 0.32);
    }
    .sidebar-toggle {
        position: absolute;
        top: 1.2rem;
        right: -18px;
        background: linear-gradient(135deg, #0057B8 0%, #FFD700 100%);
        color: #fff;
        border-radius: 50%;
        width: 36px; height: 36px;
        display: flex; align-items: center; justify-content: center;
        box-shadow: 0 2px 12px #FFD70044;
        border: none;
        cursor: pointer;
        z-index: 10002;
        transition: background 0.2s, transform 0.1s, box-shadow 0.2s cubic-bezier(.4,2,.6,1);
        animation: pulse 1.6s cubic-bezier(.4,2,.6,1) infinite;
    }
    @keyframes pulse {
        0% { box-shadow: 0 0 0 0 #FFD70044; }
        70% { box-shadow: 0 0 0 12px #FFD70011; }
        100% { box-shadow: 0 0 0 0 #FFD70044; }
    }
    .sidebar-toggle:hover { background: linear-gradient(135deg, #FFD700 0%, #0057B8 100%); }
    .sidebar-toggle:active {
        transform: scale(0.92) rotate(-8deg);
        box-shadow: 0 0 0 0 #FFD70044;
    }
    .sidebar-avatar {
        width: 48px; height: 48px;
        border-radius: 50%;
        margin-bottom: 1.5rem;
        border: 2.5px solid #FFD700;
        box-shadow: 0 2px 12px #FFD70033;
        background: url('https://randomuser.me/api/portraits/men/32.jpg') center/cover;
        transition: border 0.3s, box-shadow 0.3s;
    }
    .sidebar-section {
        display: flex;
        align-items: center;
        gap: 1.1rem;
        padding: 0.85rem 1.1rem;
        border-radius: 16px;
        font-size: 1.13rem;
        font-weight: 600;
        color: #eaf6ff;
        cursor: pointer;
        transition: background 0.18s cubic-bezier(.4,2,.6,1), color 0.18s, box-shadow 0.18s, transform 0.18s cubic-bezier(.4,2,.6,1);
        border: 1.5px solid transparent;
        width: 100%;
        margin-bottom: 0.2rem;
        position: relative;
        opacity: 0.92;
        backdrop-filter: blur(2px);
        animation: fadein 0.7s;
    }
    @keyframes fadein { from { opacity: 0; transform: translateX(-16px);} to { opacity: 0.92; transform: none; } }
    .sidebar-section.selected {
        background: rgba(0,87,184,0.22);
        color: #FFD700;
        border: 1.5px solid #FFD700;
        box-shadow: 0 2px 16px #FFD70033;
        transform: scale(1.04);
    }
    .sidebar-section:hover {
        background: rgba(0,87,184,0.18);
        color: #FFD700;
        transform: scale(1.03);
    }
    .sidebar-section:active {
        transform: scale(0.97) translateX(2px);
        box-shadow: 0 2px 16px #FFD70055;
    }
    .sidebar-section .tooltip {
        visibility: hidden;
        opacity: 0;
        background: rgba(0,0,0,0.7);
        color: #fff;
        border-radius: 8px;
        padding: 0.3rem 0.7rem;
        position: absolute;
        left: 60px;
        top: 50%;
        transform: translateY(-50%);
        font-size: 0.98rem;
        white-space: nowrap;
        z-index: 10003;
        transition: opacity 0.2s;
        animation: fadein 0.4s cubic-bezier(.4,2,.6,1);
    }
    .sidebar-section:hover .tooltip {
        visibility: visible;
        opacity: 1;
    }
    .sidebar-bottom {
        margin-top: auto;
        width: 100%;
        display: flex;
        flex-direction: column;
        gap: 1.2rem;
        align-items: flex-start;
    }
    .sidebar-settings {
        color: #FFD700;
        font-size: 1.5rem;
        cursor: pointer;
        margin-left: 0.5rem;
        margin-bottom: 0.5rem;
        transition: color 0.2s;
    }
    .sidebar-settings:hover { color: #0057B8; }
    .main-content-glass {
        margin-left: 90px;
        margin-top: 2.5rem;
        padding: 2.5rem 2.5rem 2.5rem 2.5rem;
        background: rgba(255,255,255,0.22);
        box-shadow: 0 8px 32px 0 rgba(31,38,135,0.18);
        backdrop-filter: blur(32px) saturate(180%);
        -webkit-backdrop-filter: blur(32px) saturate(180%);
        border-radius: 32px;
        border: 2px solid rgba(255,255,255,0.18);
        min-height: 80vh;
        max-width: 1100px;
        transition: box-shadow 0.22s cubic-bezier(.4,2,.6,1), background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1), opacity 0.22s cubic-bezier(.4,2,.6,1);
        will-change: box-shadow, background, transform, opacity;
        animation: fadein 0.8s;
    }
    .main-content-glass.expanded { margin-left: 280px; }
    .glass-card {
        background: rgba(255,255,255,0.22);
        box-shadow: 0 8px 32px 0 rgba(31,38,135,0.18);
        backdrop-filter: blur(24px) saturate(180%);
        -webkit-backdrop-filter: blur(24px) saturate(180%);
        border-radius: 24px;
        border: 1.5px solid rgba(255,255,255,0.18);
        padding: 2.5rem 2rem 2rem 2rem;
        margin-bottom: 2rem;
        transition: box-shadow 0.22s cubic-bezier(.4,2,.6,1), background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1), opacity 0.22s cubic-bezier(.4,2,.6,1);
        will-change: box-shadow, background, transform, opacity;
        animation: fadein 0.6s cubic-bezier(.4,2,.6,1);
    }
    .glass-card:hover {
        box-shadow: 0 12px 48px 0 rgba(31,38,135,0.22);
        transform: translateY(-2px) scale(1.01);
        background: rgba(255,255,255,0.28);
    }
    .glass-input, .glass-select, .glass-slider {
        background: rgba(255,255,255,0.7) !important;
        color: #222 !important;
        border-radius: 18px !important;
        border: 1.5px solid #e0e7ef !important;
        font-size: 1.15rem !important;
        box-shadow: 0 4px 24px rgba(0,0,0,0.04) !important;
        transition: box-shadow 0.18s cubic-bezier(.4,2,.6,1), background 0.18s cubic-bezier(.4,2,.6,1);
    }
    .glass-input:focus, .glass-select:focus, .glass-slider:focus {
        box-shadow: 0 0 0 4px #FFD70055 !important;
    }
    .glass-btn {
        box-shadow: 0 2px 16px rgba(0,87,184,0.10);
        background: linear-gradient(90deg, #e3e9f7 0%, #eaf6ff 100%);
        color: #0057B8;
        border-radius: 18px;
        font-weight: 700;
        font-size: 1.1rem;
        transition: background 0.18s cubic-bezier(.4,2,.6,1), color 0.18s, box-shadow 0.18s, transform 0.18s cubic-bezier(.4,2,.6,1);
        border: 1.5px solid #b3d4fc;
        position: relative;
        overflow: hidden;
        animation: fadein 0.7s;
        will-change: background, color, box-shadow, transform;
    }
    .glass-btn:hover {
        background: linear-gradient(90deg, #d0e6ff 0%, #b3d4fc 100%);
        color: #003366;
        transform: scale(1.045);
        box-shadow: 0 4px 32px #FFD70044;
    }
    .glass-btn:active {
        transform: scale(0.96);
        box-shadow: 0 2px 16px #FFD70077;
    }
    .glass-btn::after {
        content: '';
        position: absolute;
        left: 50%; top: 50%;
        width: 0; height: 0;
        background: rgba(0,87,184,0.18);
        border-radius: 50%;
        transform: translate(-50%, -50%);
        transition: width 0.35s cubic-bezier(.4,2,.6,1), height 0.35s cubic-bezier(.4,2,.6,1), opacity 0.35s;
        opacity: 0.5;
        pointer-events: none;
    }
    .glass-btn:active::after {
        width: 180%; height: 180%; opacity: 0;
        transition: width 0.35s cubic-bezier(.4,2,.6,1), height 0.35s cubic-bezier(.4,2,.6,1), opacity 0.35s;
    }
    @media (max-width: 900px) {
        .main-content-glass, .main-content-glass.expanded { margin-left: 0 !important; max-width: 100vw; border-radius: 0; }
        .glass-sidebar, .glass-sidebar.expanded { position: static; width: 100vw !important; flex-direction: row; align-items: center; border-radius: 0; box-shadow: none; }
    }
    </style>
''', unsafe_allow_html=True)

# Aurora/bokeh animated background layer
st.markdown('<div class="aurora-bg"></div>', unsafe_allow_html=True)

# Live theme/customization panel (in sidebar bottom)
if "glass_opacity" not in st.session_state:
    st.session_state["glass_opacity"] = 0.22
if "accent_color" not in st.session_state:
    st.session_state["accent_color"] = "#FFD700"
if "bg_style" not in st.session_state:
    st.session_state["bg_style"] = "Aurora"
if "font_family" not in st.session_state:
    st.session_state["font_family"] = "Inter"

# Glassy sticky toolbar (at top of page)
st.markdown('''
    <div class="glass-toolbar">
        <span class="toolbar-title">🧠 Obsidian Protocol v2.0</span>
        <span class="toolbar-tip">💡 Did you know? You can batch process files and visualize knowledge graphs!</span>
    </div>
''', unsafe_allow_html=True)

# Floating Action Button (FAB) for quick Analyze
def render_fab():
    st.markdown('''
        <button class="fab" onclick="window.scrollTo({top: 0, behavior: 'smooth'});">🔍</button>
    ''', unsafe_allow_html=True)
render_fab()

# Wrap main content in glass-card for luxury look
from contextlib import contextmanager
@contextmanager
def glass_card():
    st.markdown('<div class="glass-card">', unsafe_allow_html=True)
    yield
    st.markdown('</div>', unsafe_allow_html=True)

# Onboarding tour (show only once per session)
if "onboarded" not in st.session_state:
    st.session_state["onboarded"] = False
if not st.session_state["onboarded"]:
    st.info("👋 Welcome to Obsidian Protocol! Upload or paste your content, then explore the tabs for powerful AI analysis, visualizations, and more. Use the Theme Picker to personalize your experience.")
    if st.button("Got it! Start Exploring", key="onboard_btn"):
        st.session_state["onboarded"] = True

# Personalized greeting (top of main content)
user = getpass.getuser() if hasattr(getpass, 'getuser') else "User"
st.markdown(f'<div style="font-size:1.25rem;font-weight:600;margin-bottom:1.2rem;color:#FFD700;">👋 Welcome back, {user}! You are special. ✨</div>', unsafe_allow_html=True)

# Animated loader bar (example, can be shown during processing)
def show_loader():
    st.markdown('<div class="glass-loader"></div>', unsafe_allow_html=True)

# Confetti/sparkle on key actions (use st.balloons or st.snow for now, can be replaced with JS confetti)
def celebrate():
    st.balloons()
    st.success("🎉 Analysis complete! Enjoy your insights.")

# Collapsible sidebar state
if "sidebar_expanded" not in st.session_state:
    st.session_state["sidebar_expanded"] = False

# Sidebar navigation with avatar, toggle, and settings
SECTIONS = [
    ("Analyze", "🔍"),
    ("AI Tools", "🤖"),
    ("Q&A", "❓"),
    ("Batch", "📦"),
    ("Captioning", "🖼️"),
    ("Audio Analysis", "🎤"),
    ("Graph", "🕸️"),
    ("Similarity", "🧬"),
    ("Workflow", "🔗"),
    ("Timeline", "🗓️"),
    ("Topics", "📚"),
    ("Explain", "💡"),
    ("Settings", "⚙️"),
    ("History", "🕑"),
    ("Export", "⬇️"),
    ("Accessibility", "🦾"),
    ("User Profiles", "👤"),
    ("Real-Time Collaboration", "🤝"),
    ("Data Visualizations", "📊"),
    ("Theme Picker", "🎨"),
    ("Journey", "🛤️")
]
if "active_section" not in st.session_state:
    st.session_state["active_section"] = SECTIONS[0][0]

# Sidebar HTML with improved styling
sidebar_class = "glass-sidebar expanded" if st.session_state["sidebar_expanded"] else "glass-sidebar"
sidebar_html = f'<div class="{sidebar_class}">' + \
    '<div class="sidebar-avatar" title="Profile"></div>' + \
    '<button class="sidebar-toggle" onclick="window.parent.postMessage({isStreamlitMessage: true, type: \'sidebar:toggle\'}, \'*\')">' + ("⏪" if st.session_state["sidebar_expanded"] else "⏩") + '</button>'
for section, icon in SECTIONS:
    selected = "selected" if st.session_state["active_section"] == section else ""
    tooltip = f'<span class="tooltip">{section}</span>'
    sidebar_html += f'<div class="sidebar-section {selected}" onclick="window.location.hash=\"#{section}\";window.dispatchEvent(new Event(\'hashchange\'))">{icon} {tooltip if not st.session_state["sidebar_expanded"] else section}</div>'
sidebar_html += '<div class="sidebar-bottom">'
sidebar_html += '<span class="sidebar-settings" title="Settings">⚙️</span>'
sidebar_html += '<span class="sidebar-settings" title="Theme">🎨</span>'
sidebar_html += '</div></div>'
st.markdown(sidebar_html, unsafe_allow_html=True)

# JS for sidebar toggle and section switching
st.markdown('''
    <script>
    window.addEventListener('hashchange', function() {
        const section = window.location.hash.replace('#','');
        if (section) {
            window.parent.postMessage({isStreamlitMessage: true, type: 'streamlit:setComponentValue', value: section}, '*');
        }
    });
    window.addEventListener('message', function(event) {
        if (event.data && event.data.isStreamlitMessage && event.data.type === 'sidebar:toggle') {
            window.parent.postMessage({isStreamlitMessage: true, type: 'streamlit:rerun'}, '*');
        }
    });
    </script>
''', unsafe_allow_html=True)

# Main content area as a glass card, ChatGPT-style
main_class = "main-content-glass expanded" if st.session_state["sidebar_expanded"] else "main-content-glass"
st.markdown(f'<div class="{main_class}" style="padding: 0; background: transparent;">', unsafe_allow_html=True)

# Session state for results and chat
if "results" not in st.session_state:
    st.session_state["results"] = ""
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []
if "ai_outputs" not in st.session_state:
    st.session_state["ai_outputs"] = {}
if "batch_results" not in st.session_state:
    st.session_state["batch_results"] = []

# Model caching for performance
@st.cache_resource(show_spinner=False)
def get_summarizer():
    if pipeline:
        return pipeline("summarization", model="sshleifer/distilbart-cnn-12-6")
    return None
@st.cache_resource(show_spinner=False)
def get_sentiment():
    if pipeline:
        return pipeline("sentiment-analysis", model="distilbert-base-uncased-finetuned-sst-2-english")
    return None
@st.cache_resource(show_spinner=False)
def get_translator(lang="fr"):
    if pipeline:
        return pipeline(f"translation_en_to_{lang}", model=f"Helsinki-NLP/opus-mt-en-{lang}")
    return None
@st.cache_resource(show_spinner=False)
def get_qa():
    if pipeline:
        return pipeline("question-answering", model="distilbert-base-cased-distilled-squad")
    return None
@st.cache_resource(show_spinner=False)
def get_blip():
    if BlipProcessor and BlipForConditionalGeneration:
        processor = BlipProcessor.from_pretrained("Salesforce/blip-image-captioning-base")
        model = BlipForConditionalGeneration.from_pretrained("Salesforce/blip-image-captioning-base")
        return processor, model
    return None, None
@st.cache_resource(show_spinner=False)
def get_sentence_transformer():
    if SentenceTransformer:
        return SentenceTransformer('all-MiniLM-L6-v2')
    return None

# --- IMPROVED CHATGPT-STYLE GLASSY CHAT INTERFACE ---
# Enhanced sidebar that works with main navigation
st.markdown('''
    <style>
    /* Enhance existing glass-sidebar with chat-style improvements */
    .glass-sidebar {
        background: rgba(10, 10, 20, 0.75) !important;
        border-right: 2px solid rgba(255,255,255,0.12);
        box-shadow: 4px 0 32px #000A1A66;
        backdrop-filter: blur(20px) !important;
        transition: all 0.3s cubic-bezier(.4,2,.6,1);
    }
    .glass-sidebar .sidebar-avatar {
        width: 48px; height: 48px;
        border-radius: 50%;
        border: 2.5px solid #FFD700;
        background: url('https://randomuser.me/api/portraits/men/32.jpg') center/cover;
        margin: 1rem auto 2rem auto;
        transition: border 0.3s, box-shadow 0.3s;
        box-shadow: 0 2px 12px #FFD70033;
        display: block;
    }
    .glass-sidebar .sidebar-avatar:hover {
        box-shadow: 0 0 0 8px #FFD70033, 0 2px 12px #FFD70033;
        border: 2.5px solid #FFF;
        animation: pulse 1.2s infinite;
    }
    .sidebar-section {
        transition: all 0.18s cubic-bezier(.4,2,.6,1) !important;
        border-radius: 12px !important;
        margin: 0.5rem !important;
        padding: 0.8rem !important;
    }
    .sidebar-section:hover {
        background: rgba(255,215,0,0.15) !important;
        transform: translateX(4px) !important;
    }
    .sidebar-section.selected {
        background: rgba(255,215,0,0.25) !important;
        border-left: 3px solid #FFD700 !important;
    }
    </style>
''', unsafe_allow_html=True)

# Add avatar to existing sidebar
sidebar_avatar_html = '<div class="sidebar-avatar" title="Profile"></div>'

# --- Enhanced Chat Interface for Q&A Section ---
st.markdown('''
    <style>
    .enhanced-chat-container {
        padding: 2rem;
        background: rgba(255,255,255,0.05);
        box-shadow: 0 8px 32px 0 rgba(0,0,0,0.3);
        backdrop-filter: blur(20px) saturate(180%);
        -webkit-backdrop-filter: blur(20px) saturate(180%);
        border-radius: 24px;
        border: 1px solid rgba(255,255,255,0.1);
        margin: 1rem 0;
        animation: fadein 0.8s cubic-bezier(.4,2,.6,1);
        position: relative;
        max-width: 1000px;
        margin-left: auto;
        margin-right: auto;
    }
    .chat-bubble {
        max-width: 80%;
        margin-bottom: 1.2rem;
        padding: 1.2rem 1.8rem;
        border-radius: 18px;
        color: #FFFFFF;
        font-size: 1.1rem;
        line-height: 1.6;
        box-shadow: 0 4px 12px rgba(0,0,0,0.2);
        backdrop-filter: blur(10px);
        animation: fadein 0.5s cubic-bezier(.4,2,.6,1);
        position: relative;
        transition: all 0.3s ease;
        word-wrap: break-word;
    }
    .chat-bubble.user {
        background: rgba(255,215,0,0.15);
        color: #FFD700;
        margin-left: auto;
        border: 1px solid rgba(255,215,0,0.3);
    }
    .chat-bubble.ai {
        background: rgba(255,255,255,0.08);
        color: #FFFFFF;
        margin-right: auto;
        border: 1px solid rgba(255,255,255,0.1);
    }
    .chat-bubble:hover {
        transform: translateY(-2px);
        box-shadow: 0 6px 20px rgba(0,0,0,0.3);
    }
    /* Removed problematic chat-input-row styling */
    .stTextInput > div > div > input {
        font-size: 1.1rem !important;
        padding: 1rem 1.5rem !important;
        border-radius: 25px !important;
        border: 2px solid rgba(255,215,0,0.3) !important;
        background: rgba(255,255,255,0.1) !important;
        color: #FFFFFF !important;
        box-shadow: 0 4px 15px rgba(0,0,0,0.2) !important;
        outline: none !important;
        transition: all 0.3s ease !important;
        backdrop-filter: blur(15px) !important;
        width: 100% !important;
        height: 55px !important;
    }
    .stTextInput > div > div > input:focus {
        border: 2px solid #FFD700 !important;
        box-shadow: 0 0 0 3px rgba(255,215,0,0.3), 0 6px 25px rgba(0,0,0,0.3) !important;
        background: rgba(255,255,255,0.15) !important;
        transform: translateY(-2px);
    }
    /* Cleaned up unused input animations */
    .chat-send-btn, .stButton > button {
        font-size: 1.1rem !important;
        font-weight: 600 !important;
        padding: 1rem 2rem !important;
        border-radius: 24px !important;
        border: none !important;
        background: linear-gradient(135deg, #FFD700 0%, #FFA500 100%) !important;
        color: #000000 !important;
        box-shadow: 0 4px 20px rgba(255,215,0,0.3) !important;
        cursor: pointer !important;
        transition: all 0.3s cubic-bezier(.4,2,.6,1) !important;
        position: relative !important;
        overflow: hidden !important;
        height: auto !important;
        min-height: 50px !important;
    }
    .chat-send-btn:hover, .stButton > button:hover {
        background: linear-gradient(135deg, #FFA500 0%, #FFD700 100%) !important;
        box-shadow: 0 6px 25px rgba(255,215,0,0.4) !important;
        transform: translateY(-2px) !important;
    }
    .chat-send-btn:active, .stButton > button:active {
        transform: scale(0.96) !important;
        box-shadow: 0 2px 15px rgba(255,215,0,0.5) !important;
    }
    .chat-send-btn::after {
        content: '';
        position: absolute;
        left: 50%; top: 50%;
        width: 0; height: 0;
        background: rgba(0,87,184,0.18);
        border-radius: 50%;
        transform: translate(-50%, -50%);
        transition: width 0.35s cubic-bezier(.4,2,.6,1), height 0.35s cubic-bezier(.4,2,.6,1), opacity 0.35s;
        opacity: 0.5;
        pointer-events: none;
    }
    .chat-send-btn:active::after {
        width: 180%; height: 180%; opacity: 0;
        transition: width 0.35s cubic-bezier(.4,2,.6,1), height 0.35s cubic-bezier(.4,2,.6,1), opacity 0.35s;
    }
    .chat-input-icons {
        display: flex;
        gap: 0.8rem;
        align-items: center;
    }
    .chat-input-icon {
        font-size: 1.5rem;
        color: #FFD700;
        background: rgba(0,0,0,0.12);
        border-radius: 50%;
        padding: 0.4rem;
        box-shadow: 0 2px 8px #FFD70022;
        cursor: pointer;
        transition: background 0.18s, color 0.18s, transform 0.18s;
    }
    .chat-input-icon.selected, .chat-input-icon:hover {
        background: #FFD700;
        color: #232946;
        transform: scale(1.12);
    }
    .chat-preview {
        margin-top: 0.7rem;
        margin-bottom: 0.7rem;
        background: rgba(255,255,255,0.13);
        border-radius: 16px;
        padding: 1rem 1.2rem;
        color: #FFD700;
        font-size: 1.05rem;
        box-shadow: 0 2px 8px #FFD70022;
        animation: fadein 0.5s cubic-bezier(.4,2,.6,1);
    }
    .chat-suggestions {
        display: flex;
        gap: 0.7rem;
        margin-top: 0.5rem;
        margin-bottom: 0.5rem;
        flex-wrap: wrap;
    }
    .chat-suggestion {
        background: rgba(0,87,184,0.18);
        color: #FFD700;
        border-radius: 12px;
        padding: 0.4rem 1.1rem;
        font-size: 1.01rem;
        cursor: pointer;
        transition: background 0.18s, color 0.18s, transform 0.18s;
        box-shadow: 0 2px 8px #FFD70022;
        animation: fadein 0.4s cubic-bezier(.4,2,.6,1);
    }
    .chat-suggestion:hover {
        background: #FFD700;
        color: #232946;
        transform: scale(1.08);
    }
    
    /* Style file uploaders */
    .stFileUploader > div {
        background: rgba(255,255,255,0.08) !important;
        border: 2px dashed rgba(255,215,0,0.4) !important;
        border-radius: 16px !important;
        padding: 1.5rem !important;
        transition: all 0.3s ease !important;
        text-align: center !important;
    }
    .stFileUploader > div:hover {
        border-color: #FFD700 !important;
        background: rgba(255,215,0,0.1) !important;
        transform: translateY(-2px) !important;
        box-shadow: 0 6px 20px rgba(0,0,0,0.2) !important;
    }
    .stFileUploader label {
        color: #FFD700 !important;
        font-weight: 600 !important;
        font-size: 1rem !important;
    }
    .stFileUploader div[data-testid="stFileUploaderDropzoneInstructions"] {
        color: rgba(255,255,255,0.7) !important;
    }
    </style>
''', unsafe_allow_html=True)

# --- SMART PROCESSING FUNCTIONS ---
def process_chat_input(user_input, context=""):
    """Process user input and return appropriate AI response"""
    input_lower = user_input.lower()
    
    # Summarization requests
    if any(word in input_lower for word in ["summarize", "summary", "tldr", "brief"]):
        if context:
            summarizer = get_summarizer()
            if summarizer:
                try:
                    result = summarizer(context[:1024], max_length=150, min_length=50, do_sample=False)
                    return f"📝 **Summary:**\n\n{result[0]['summary_text']}"
                except:
                    return f"📝 **Quick Summary:**\n\n{context[:200]}...\n\n*Note: Advanced summarization unavailable - showing preview instead.*"
            else:
                return f"📝 **Quick Summary:**\n\n{context[:300]}...\n\n*Advanced summarization requires additional libraries.*"
        else:
            return "📝 I'd be happy to summarize content for you! Please upload a document or provide some text first."
    
    # Translation requests
    elif any(word in input_lower for word in ["translate", "french", "spanish", "german"]):
        if context:
            return f"🌍 **Translation:**\n\nI can translate your content! Please specify the target language (French, Spanish, German, etc.)\n\n**Content to translate:**\n{context[:200]}..."
        else:
            return "🌍 I can translate content for you! Please upload a document first, then tell me which language you'd like."
    
    # Analysis requests
    elif any(word in input_lower for word in ["analyze", "analysis", "insights", "extract"]):
        if context:
            return f"🔍 **Analysis Results:**\n\n**Content Length:** {len(context)} characters\n**Word Count:** ~{len(context.split())} words\n\n**Key Insights:**\n• Content appears to be {detect_content_type(context)}\n• Main topics detected\n• Ready for deeper analysis\n\n**Available Analysis:**\n- Sentiment analysis\n- Entity extraction\n- Topic modeling\n- Question answering"
        else:
            return "🔍 I'm ready to analyze content for you! Please upload a document, image, or audio file first."
    
    # Question answering
    elif any(word in input_lower for word in ["what", "how", "why", "when", "where", "who", "?"]):
        if context:
            qa_pipeline = get_qa()
            if qa_pipeline:
                try:
                    result = qa_pipeline(question=user_input, context=context[:1000])
                    return f"❓ **Answer:**\n\n{result['answer']}\n\n*Confidence: {result['score']:.2%}*"
                except:
                    return f"❓ **Answer:**\n\nBased on the uploaded content, I can see information related to your question. However, I need more specific context or the advanced Q&A model to provide a precise answer."
            else:
                return f"❓ **Answer:**\n\nI can help answer questions about your uploaded content! The Q&A feature requires additional libraries for optimal performance."
        else:
            return "❓ I can answer questions about uploaded content! Please upload a document first, then ask your question."
    
    # Sentiment analysis
    elif any(word in input_lower for word in ["sentiment", "emotion", "feeling", "mood"]):
        if context:
            sentiment_analyzer = get_sentiment()
            if sentiment_analyzer:
                try:
                    result = sentiment_analyzer(context[:512])
                    sentiment = result[0]['label']
                    confidence = result[0]['score']
                    return f"😊 **Sentiment Analysis:**\n\n**Overall Sentiment:** {sentiment}\n**Confidence:** {confidence:.2%}\n\n**Analysis:** The content appears to be {sentiment.lower()} in tone."
                except:
                    return f"😊 **Sentiment Analysis:**\n\nI can analyze the sentiment of your content! The advanced sentiment analysis requires additional libraries."
            else:
                return f"😊 **Sentiment Analysis:**\n\nI can analyze sentiment for you! This feature requires additional libraries to be installed."
        else:
            return "😊 I can analyze the sentiment of content! Please upload some text first."
    
    # Default helpful response
    else:
        return f"💡 **I can help you with:**\n\n• **Analyze** documents, images, audio\n• **Summarize** long content\n• **Translate** to different languages\n• **Answer questions** about your content\n• **Extract insights** and key information\n• **Sentiment analysis** of text\n\n**Your message:** '{user_input}'\n\nTry uploading a file or asking a specific question!"

def detect_content_type(content):
    """Detect the type of content"""
    if len(content) < 100:
        return "short text or notes"
    elif any(word in content.lower() for word in ["abstract", "introduction", "conclusion", "references"]):
        return "academic or research document"
    elif any(word in content.lower() for word in ["article", "news", "report"]):
        return "news article or report"
    elif content.count('\n') > 20:
        return "structured document or data"
    else:
        return "general text content"

def process_uploaded_file(uploaded_file):
    """Process uploaded file and extract content"""
    try:
        if uploaded_file.name.endswith('.txt'):
            content = uploaded_file.read().decode("utf-8")
            return content
        elif uploaded_file.name.endswith('.pdf') and PyPDF2:
            reader = PyPDF2.PdfReader(uploaded_file)
            content = "\n".join([page.extract_text() or "" for page in reader.pages])
            return content
        else:
            return None
    except Exception as e:
        return None

def process_uploaded_image(uploaded_image):
    """Process uploaded image and extract text"""
    try:
        if Image and pytesseract:
            img = Image.open(uploaded_image)
            content = pytesseract.image_to_string(img)
            return content.strip() if content.strip() else None
        else:
            return None
    except Exception as e:
        return None

def process_uploaded_audio(uploaded_audio):
    """Process uploaded audio and transcribe"""
    try:
        if whisper:
            with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(uploaded_audio.name)[1]) as tmp_file:
                tmp_file.write(uploaded_audio.read())
                tmp_path = tmp_file.name
            
            model = whisper.load_model("base")
            result = model.transcribe(tmp_path)
            content = result["text"]
            os.remove(tmp_path)
            return content
        else:
            return None
    except Exception as e:
        return None

def process_url(url):
    """Process URL and extract content"""
    try:
        if requests and BeautifulSoup:
            r = requests.get(url, timeout=10)
            soup = BeautifulSoup(r.text, "html.parser")
            paragraphs = [p.get_text() for p in soup.find_all("p")]
            content = "\n".join(paragraphs)
            return content.strip() if content.strip() else None
        else:
            return None
    except Exception as e:
        return None

# --- MAIN CHATGPT-STYLE INTERFACE ---
# Chat state initialization
if "chat_history" not in st.session_state:
    st.session_state["chat_history"] = []
if "chat_input" not in st.session_state:
    st.session_state["chat_input"] = ""

# Main ChatGPT-style container
st.markdown('<div class="enhanced-chat-container">', unsafe_allow_html=True)

# Header with Obsidian Protocol branding and controls
header_col1, header_col2, header_col3 = st.columns([2, 6, 2])

with header_col1:
    if st.button("🗑️ Clear Chat", help="Clear conversation history"):
        st.session_state["chat_history"] = []
        st.session_state["results"] = ""
        st.rerun()

with header_col2:
    st.markdown('''
        <div style="text-align: center;">
            <h1 style="color: #FFD700; font-size: 2.5rem; font-weight: 700; text-shadow: 0 4px 16px #FFD70044; margin-bottom: 0.5rem;">
                Obsidian Protocol
            </h1>
            <p style="color: rgba(255,255,255,0.7); font-size: 1.1rem; margin: 0;">
                AI-Powered Analysis & Intelligence Platform
            </p>
        </div>
    ''', unsafe_allow_html=True)

with header_col3:
    st.write(f"💬 {len(st.session_state['chat_history'])} messages")

# Chat history display
chat_history_container = st.container()
with chat_history_container:
    if st.session_state["chat_history"]:
        for i, (sender, message) in enumerate(st.session_state["chat_history"]):
            bubble_class = "chat-bubble user" if sender == "user" else "chat-bubble ai"
            st.markdown(f'<div class="{bubble_class}">{message}</div>', unsafe_allow_html=True)
    else:
        # Welcome message
        st.markdown('''
            <div class="chat-bubble ai" style="margin: 2rem auto; max-width: 80%; text-align: center;">
                👋 Welcome to Obsidian Protocol! I can help you analyze files, extract insights, answer questions, and much more. 
                <br><br>
                Try uploading a document, asking a question, or selecting from the suggestions below.
            </div>
        ''', unsafe_allow_html=True)

# Interactive suggestions
suggestion_clicked = False
col1, col2, col3, col4 = st.columns(4)

with col1:
    if st.button("📄 Analyze Document", use_container_width=True):
        st.session_state["chat_history"].append(("user", "Analyze this document"))
        ai_response = process_chat_input("Analyze this document", st.session_state.get("results", ""))
        st.session_state["chat_history"].append(("ai", ai_response))
        st.rerun()

with col2:
    if st.button("📝 Summarize", use_container_width=True):
        st.session_state["chat_history"].append(("user", "Summarize this content"))
        ai_response = process_chat_input("Summarize this content", st.session_state.get("results", ""))
        st.session_state["chat_history"].append(("ai", ai_response))
        st.rerun()

with col3:
    if st.button("💡 Extract Insights", use_container_width=True):
        st.session_state["chat_history"].append(("user", "Extract key insights"))
        ai_response = process_chat_input("Extract key insights", st.session_state.get("results", ""))
        st.session_state["chat_history"].append(("ai", ai_response))
        st.rerun()

with col4:
    if st.button("😊 Sentiment Analysis", use_container_width=True):
        st.session_state["chat_history"].append(("user", "Analyze sentiment"))
        ai_response = process_chat_input("Analyze sentiment", st.session_state.get("results", ""))
        st.session_state["chat_history"].append(("ai", ai_response))
        st.rerun()

# Main chat input area
chat_input_col1, chat_input_col2 = st.columns([10, 1])

with chat_input_col1:
    user_input = st.text_input(
        "",
        placeholder="Type your message, upload a file, or ask a question...",
        key="main_chat_input",
        label_visibility="collapsed"
    )

with chat_input_col2:
    send_clicked = st.button("Send", type="primary", use_container_width=True)

# File upload options
st.markdown("---")
upload_col1, upload_col2, upload_col3, upload_col4 = st.columns(4)

with upload_col1:
    uploaded_file = st.file_uploader("📄 Upload Document", type=["txt", "pdf", "docx"], label_visibility="collapsed")

with upload_col2:
    uploaded_image = st.file_uploader("🖼️ Upload Image", type=["jpg", "jpeg", "png"], label_visibility="collapsed")

with upload_col3:
    uploaded_audio = st.file_uploader("🎤 Upload Audio", type=["mp3", "wav", "m4a"], label_visibility="collapsed")

with upload_col4:
    url_input = st.text_input("🔗 Enter URL", placeholder="https://example.com", label_visibility="collapsed")

# Process input and file uploads
if send_clicked and user_input:
    # Add user message to chat
    st.session_state["chat_history"].append(("user", user_input))
    
    # Smart AI processing based on user input
    ai_response = process_chat_input(user_input, st.session_state.get("results", ""))
    st.session_state["chat_history"].append(("ai", ai_response))
    
    # Clear input
    st.session_state["chat_input"] = ""
    st.rerun()

# Handle file uploads with actual processing
if uploaded_file:
    st.session_state["chat_history"].append(("user", f"📄 Uploaded file: {uploaded_file.name}"))
    
    # Process the file
    content = process_uploaded_file(uploaded_file)
    if content:
        st.session_state["results"] = content
        ai_response = f"✅ Successfully processed '{uploaded_file.name}'!\n\n**Content Preview:**\n{content[:300]}{'...' if len(content) > 300 else ''}\n\n💡 I can now help you:\n- Summarize the content\n- Extract key insights\n- Answer questions about it\n- Translate it\n- Create visualizations"
    else:
        ai_response = f"❌ I had trouble processing '{uploaded_file.name}'. Please make sure it's a valid text, PDF, or document file."
    
    st.session_state["chat_history"].append(("ai", ai_response))
    st.rerun()

if uploaded_image:
    st.session_state["chat_history"].append(("user", f"🖼️ Uploaded image: {uploaded_image.name}"))
    
    # Process the image
    content = process_uploaded_image(uploaded_image)
    if content:
        st.session_state["results"] = content
        ai_response = f"✅ Successfully analyzed '{uploaded_image.name}'!\n\n**Extracted Text:**\n{content[:300]}{'...' if len(content) > 300 else ''}\n\n💡 I can now help you analyze this content further."
    else:
        ai_response = f"❌ I couldn't extract text from '{uploaded_image.name}'. The image might not contain readable text or the OCR libraries aren't available."
    
    st.session_state["chat_history"].append(("ai", ai_response))
    st.rerun()

if uploaded_audio:
    st.session_state["chat_history"].append(("user", f"🎤 Uploaded audio: {uploaded_audio.name}"))
    
    # Process the audio
    content = process_uploaded_audio(uploaded_audio)
    if content:
        st.session_state["results"] = content
        ai_response = f"✅ Successfully transcribed '{uploaded_audio.name}'!\n\n**Transcript:**\n{content[:300]}{'...' if len(content) > 300 else ''}\n\n💡 I can now help you analyze this transcript."
    else:
        ai_response = f"❌ I couldn't transcribe '{uploaded_audio.name}'. The Whisper library might not be available or the audio format isn't supported."
    
    st.session_state["chat_history"].append(("ai", ai_response))
    st.rerun()

if url_input:
    st.session_state["chat_history"].append(("user", f"🔗 Provided URL: {url_input}"))
    
    # Process the URL
    content = process_url(url_input)
    if content:
        st.session_state["results"] = content
        ai_response = f"✅ Successfully scraped content from: {url_input}\n\n**Content Preview:**\n{content[:300]}{'...' if len(content) > 300 else ''}\n\n💡 I can now help you analyze this web content."
    else:
        ai_response = f"❌ I couldn't scrape content from: {url_input}. The URL might be invalid or the required libraries aren't available."
    
    st.session_state["chat_history"].append(("ai", ai_response))
    st.rerun()

st.markdown('</div>', unsafe_allow_html=True)

# Close main content div
st.markdown('</div>', unsafe_allow_html=True)

# --- SECTION-BASED TOOLS (Optional - Accessible via sidebar) ---
# These sections are now secondary to the main chat interface
if st.session_state["active_section"] == "Analyze":
    st.header("Analyze File or Content")
    analyze_type = st.selectbox("Select input type:", ["Text", "TXT File", "PDF", "URL", "YouTube Video", "Image", "Audio"])
    content = ""
    lang = st.selectbox("Select language for translation/output:", ["en", "fr", "es", "de", "zh", "ar", "ru", "hi", "pt", "it"])
    if analyze_type == "Text":
        content = st.text_area("Enter text to analyze:")
    elif analyze_type == "TXT File":
        txt_file = st.file_uploader("Upload a TXT file", type=["txt"])
        if txt_file:
            content = txt_file.read().decode("utf-8")
            st.code(content, language="text")
    elif analyze_type == "PDF":
        if PyPDF2:
            pdf_file = st.file_uploader("Upload a PDF file", type=["pdf"])
            if pdf_file:
                reader = PyPDF2.PdfReader(pdf_file)
                content = "\n".join([page.extract_text() or "" for page in reader.pages])
                st.code(content, language="text")
        else:
            st.warning("PyPDF2 not installed. Run: pip install PyPDF2")
    elif analyze_type == "URL":
        if requests and BeautifulSoup:
            url = st.text_input("Enter a URL to scrape article text:")
            if url:
                try:
                    r = requests.get(url)
                    soup = BeautifulSoup(r.text, "html.parser")
                    paragraphs = [p.get_text() for p in soup.find_all("p")]
                    content = "\n".join(paragraphs)
                    st.code(content, language="text")
                except Exception as e:
                    st.error(f"Failed to fetch URL: {e}")
        else:
            st.warning("requests and beautifulsoup4 not installed. Run: pip install requests beautifulsoup4")
    elif analyze_type == "YouTube Video":
        if YouTubeTranscriptApi:
            yt_url = st.text_input("Enter YouTube video URL:")
            if yt_url:
                import re
                match = re.search(r"v=([\w-]+)", yt_url)
                if match:
                    video_id = match.group(1)
                    try:
                        transcript = YouTubeTranscriptApi.get_transcript(video_id)
                        content = " ".join([x['text'] for x in transcript])
                        st.code(content, language="text")
                    except Exception as e:
                        st.error(f"Could not fetch transcript: {e}")
                else:
                    st.warning("Invalid YouTube URL.")
        else:
            st.warning("youtube-transcript-api not installed. Run: pip install youtube-transcript-api")
    elif analyze_type == "Image":
        if Image and pytesseract:
            img_file = st.file_uploader("Upload an image", type=["jpg", "jpeg", "png"])
            if img_file:
                img = Image.open(img_file)
                st.image(img, caption="Uploaded Image", use_column_width=True)
                content = pytesseract.image_to_string(img)
                st.code(content, language="text")
        else:
            st.warning("Pillow and pytesseract not installed. Run: pip install pillow pytesseract")
    elif analyze_type == "Audio":
        if whisper:
            audio_file = st.file_uploader("Upload an audio file", type=["mp3", "wav", "m4a"])
            if audio_file:
                with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(audio_file.name)[1]) as tmp_file:
                    tmp_file.write(audio_file.read())
                    tmp_path = tmp_file.name
                model = whisper.load_model("base")
                result = model.transcribe(tmp_path)
                content = result["text"]
                st.code(content, language="text")
                os.remove(tmp_path)
        else:
            st.warning("openai-whisper not installed. Run: pip install openai-whisper")
    # Store content for further analysis
    st.session_state["results"] = content
    st.markdown("---")
    st.subheader("AI Tools (see next tab)")
    if content:
        st.download_button("Download Results as .txt", content, file_name="obsidian_output.txt")
        # TTS
        st.markdown("**Text-to-Speech (TTS):**")
        tts_engine = st.radio("TTS Engine", ["gTTS (Google)", "pyttsx3 (Offline)"])
        if st.button("Speak/Download Audio"):
            if tts_engine == "gTTS (Google)" and gTTS:
                tts = gTTS(content, lang=lang)
                tts_fp = BytesIO()
                tts.write_to_fp(tts_fp)
                tts_fp.seek(0)
                b64 = base64.b64encode(tts_fp.read()).decode()
                st.audio(f"data:audio/mp3;base64,{b64}", format="audio/mp3")
                st.download_button("Download Audio", data=tts_fp, file_name="tts.mp3")
            elif tts_engine == "pyttsx3 (Offline)" and pyttsx3:
                engine = pyttsx3.init()
                engine.save_to_file(content, "tts.wav")
                engine.runAndWait()
                with open("tts.wav", "rb") as f:
                    st.audio(f.read(), format="audio/wav")
                st.download_button("Download Audio", data=open("tts.wav", "rb"), file_name="tts.wav")
                os.remove("tts.wav")
            else:
                st.warning("TTS engine not installed. Run: pip install gtts pyttsx3")

# --- AI Tools Tab ---
if st.session_state["active_section"] == "AI Tools":
    st.header("AI Tools")
    content = st.session_state["results"]
    if not content:
        st.info("No content loaded. Please use the Analyze tab first.")
    else:
        col1, col2 = st.columns(2)
        with col1:
            st.subheader("Summarization")
            summarizer = get_summarizer()
            if summarizer:
                if st.button("Summarize", key="summarize_btn") or st.session_state["ai_outputs"].get("summary"):
                    if not st.session_state["ai_outputs"].get("summary"):
                        summary = summarizer(content, max_length=130, min_length=30, do_sample=False)[0]["summary_text"]
                        st.session_state["ai_outputs"]["summary"] = summary
                        celebrate()
                    st.success(st.session_state["ai_outputs"]["summary"])
            else:
                st.warning("Transformers not installed. Run: pip install transformers torch")
        with col2:
            st.subheader("Sentiment Analysis")
            sentiment = get_sentiment()
            if sentiment:
                if st.button("Analyze Sentiment", key="sentiment_btn") or st.session_state["ai_outputs"].get("sentiment"):
                    if not st.session_state["ai_outputs"].get("sentiment"):
                        sent = sentiment(content)[0]
                        st.session_state["ai_outputs"]["sentiment"] = f"{sent['label']} (score: {sent['score']:.2f})"
                        celebrate()
                    st.success(st.session_state["ai_outputs"]["sentiment"])
            else:
                st.warning("Transformers not installed. Run: pip install transformers torch")
        st.markdown("---")
        col3, col4 = st.columns(2)
        with col3:
            st.subheader("Named Entity Recognition (NER)")
            if nlp:
                if st.button("Extract Entities", key="ner_btn") or st.session_state["ai_outputs"].get("ner"):
                    if not st.session_state["ai_outputs"].get("ner"):
                        doc = nlp(content)
                        ents = [(ent.text, ent.label_) for ent in doc.ents]
                        st.session_state["ai_outputs"]["ner"] = ents
                        celebrate()
                    ents = st.session_state["ai_outputs"]["ner"]
                    if ents:
                        for ent, label in ents:
                            st.markdown(f"- **{ent}**: {label}")
                    else:
                        st.info("No named entities found.")
            else:
                st.warning("spaCy not installed or model missing. Run: pip install spacy && python -m spacy download en_core_web_sm")
        with col4:
            st.subheader(f"Translation (EN → {lang.upper()})")
            if lang == "en":
                translator = None
                translation_error = "Source and target language are both English—no translation needed."
            else:
                try:
                    translator = get_translator(lang)
                    translation_error = None
                except Exception as e:
                    translator = None
                    translation_error = str(e)
            if lang == "en":
                st.info("Source and target language are both English—no translation needed.")
            elif translator:
                if st.button("Translate", key="translate_btn") or st.session_state["ai_outputs"].get("translation"):
                    if not st.session_state["ai_outputs"].get("translation"):
                        try:
                            translation = translator(content)[0]['translation_text']
                            st.session_state["ai_outputs"]["translation"] = translation
                            celebrate()
                        except Exception as e:
                            st.session_state["ai_outputs"]["translation"] = f"[Translation Error]: {e}"
                    st.success(st.session_state["ai_outputs"]["translation"])
            else:
                if translation_error:
                    st.warning(f"Translation model could not be loaded: {translation_error}")
                else:
                    st.warning("Transformers not installed. Run: pip install transformers torch")

        # --- Analyze All Button ---
        st.markdown("---")
        if st.button("Analyze All", key="analyze_all_btn"):
            # Summarization
            if summarizer:
                try:
                    summary = summarizer(content, max_length=130, min_length=30, do_sample=False)[0]["summary_text"]
                except Exception as e:
                    summary = f"[Summarization Error]: {e}"
            else:
                summary = "[Summarizer not available]"
            st.session_state["ai_outputs"]["summary"] = summary
            # Sentiment
            if sentiment:
                try:
                    sent = sentiment(content)[0]
                    sentiment_result = f"{sent['label']} (score: {sent['score']:.2f})"
                except Exception as e:
                    sentiment_result = f"[Sentiment Error]: {e}"
            else:
                sentiment_result = "[Sentiment model not available]"
            st.session_state["ai_outputs"]["sentiment"] = sentiment_result
            # NER
            if nlp:
                try:
                    doc = nlp(content)
                    ents = [(ent.text, ent.label_) for ent in doc.ents]
                except Exception as e:
                    ents = [(f"[NER Error]: {e}", "ERROR")]
            else:
                ents = []
            st.session_state["ai_outputs"]["ner"] = ents
            # Translation
            if lang == "en":
                translation = "Source and target language are both English—no translation needed."
            elif translator:
                try:
                    translation = translator(content)[0]['translation_text']
                except Exception as e:
                    translation = f"[Translation Error]: {e}"
            else:
                translation = "[Translation model not available]"
            st.session_state["ai_outputs"]["translation"] = translation
            celebrate()

        # --- Display All Results if available ---
        if st.session_state["ai_outputs"].get("summary"):
            st.markdown("**Summary:**")
            st.success(st.session_state["ai_outputs"]["summary"])
        if st.session_state["ai_outputs"].get("sentiment"):
            st.markdown("**Sentiment:**")
            st.success(st.session_state["ai_outputs"]["sentiment"])
        if st.session_state["ai_outputs"].get("ner"):
            st.markdown("**Named Entities:**")
            ents = st.session_state["ai_outputs"]["ner"]
            if ents:
                for ent, label in ents:
                    st.markdown(f"- **{ent}**: {label}")
            else:
                st.info("No named entities found.")
        if st.session_state["ai_outputs"].get("translation"):
            st.markdown(f"**Translation (EN → {lang.upper()}):**")
            st.success(st.session_state["ai_outputs"]["translation"])

# --- Q&A Tab ---
if st.session_state["active_section"] == "Q&A":
    st.header("Document Q&A")
    content = st.session_state["results"]
    if not content:
        st.info("No content loaded. Please use the Analyze tab first.")
    else:
        qa = get_qa()
        question = st.text_input("Ask a question about the document:")
        if st.button("Get Answer") and question and qa:
            answer = qa(question=question, context=content)
            st.success(f"**Answer:** {answer['answer']}")

# --- Batch Tab ---
if st.session_state["active_section"] == "Batch":
    st.header("Batch Processing")
    batch_files = st.file_uploader("Upload multiple files (TXT, PDF, Image, Audio)", type=["txt", "pdf", "jpg", "jpeg", "png", "mp3", "wav", "m4a"], accept_multiple_files=True)
    batch_results = []
    if batch_files:
        for f in batch_files:
            ext = os.path.splitext(f.name)[1].lower()
            text = ""
            if ext == ".txt":
                text = f.read().decode("utf-8")
            elif ext == ".pdf" and PyPDF2:
                reader = PyPDF2.PdfReader(f)
                text = "\n".join([page.extract_text() or "" for page in reader.pages])
            elif ext in [".jpg", ".jpeg", ".png"] and Image and pytesseract:
                img = Image.open(f)
                text = pytesseract.image_to_string(img)
            elif ext in [".mp3", ".wav", ".m4a"] and whisper:
                with tempfile.NamedTemporaryFile(delete=False, suffix=ext) as tmp_file:
                    tmp_file.write(f.read())
                    tmp_path = tmp_file.name
                model = whisper.load_model("base")
                result = model.transcribe(tmp_path)
                text = result["text"]
                os.remove(tmp_path)
            batch_results.append((f.name, text))
        st.session_state["batch_results"] = batch_results
        for fname, text in batch_results:
            st.markdown(f"**{fname}:**")
            st.code(text, language="text")
        if batch_results:
            all_text = "\n\n".join([f"{fname}:\n{text}" for fname, text in batch_results])
            st.download_button("Download All Results as .txt", all_text, file_name="obsidian_batch_output.txt")

# --- Captioning Tab ---
if st.session_state["active_section"] == "Captioning":
    st.header("Image Captioning")
    if BlipProcessor and BlipForConditionalGeneration:
        img_file = st.file_uploader("Upload an image for captioning", type=["jpg", "jpeg", "png"], key="caption_img")
        if img_file:
            img = Image.open(img_file).convert("RGB")
            st.image(img, caption="Uploaded Image", use_column_width=True)
            processor, model = get_blip()
            import torch
            inputs = processor(img, return_tensors="pt")
            out = model.generate(**inputs)
            caption = processor.decode(out[0], skip_special_tokens=True)
            st.success(f"**Caption:** {caption}")
    else:
        st.warning("BLIP not installed. Run: pip install transformers torch")

# --- Audio Analysis Tab ---
if st.session_state["active_section"] == "Audio Analysis":
    st.header("Audio Sentiment & Speaker Diarization")
    audio_file = st.file_uploader("Upload audio for analysis", type=["mp3", "wav", "m4a"], key="audio_analysis")
    if audio_file and aS:
        with tempfile.NamedTemporaryFile(delete=False, suffix=os.path.splitext(audio_file.name)[1]) as tmp_file:
            tmp_file.write(audio_file.read())
            tmp_path = tmp_file.name
        st.info("Speaker diarization and audio sentiment coming soon (requires pyAudioAnalysis)")
        # Example: diarization = aS.speaker_diarization(tmp_path, n_speakers=2)
        os.remove(tmp_path)
    elif audio_file:
        st.warning("pyAudioAnalysis not installed. Run: pip install pyAudioAnalysis")

# --- Knowledge Graph Tab ---
if st.session_state["active_section"] == "Graph":
    st.header("Interactive Knowledge Graph Extraction")
    content = st.session_state["results"]
    if not content:
        st.info("No content loaded. Please use the Analyze tab first.")
    elif nx and Network and nlp:
        doc = nlp(content)
        G = nx.Graph()
        for ent in doc.ents:
            G.add_node(ent.text, label=ent.label_)
        for i in range(len(doc.ents)-1):
            G.add_edge(doc.ents[i].text, doc.ents[i+1].text)
        net = Network(height="400px", width="100%", bgcolor="#222", font_color="#FFD700")
        net.from_nx(G)
        net.save_graph("graph.html")
        with open("graph.html", "r", encoding="utf-8") as f:
            html = f.read()
        st.components.v1.html(html, height=450)
        os.remove("graph.html")
        # Clickable entities (show sentences containing entity)
        st.markdown("**Click an entity below to see related sentences:**")
        selected_entity = st.selectbox("Entities:", [ent.text for ent in doc.ents]) if doc.ents else None
        if selected_entity:
            sentences = [sent.text for sent in doc.sents if selected_entity in sent.text]
            st.markdown(f"**Sentences mentioning {selected_entity}:**")
            for s in sentences:
                st.write(s)
    else:
        st.warning("networkx, pyvis, or spaCy not installed. Run: pip install networkx pyvis spacy && python -m spacy download en_core_web_sm")

# --- Similarity/Plagiarism Tab ---
if st.session_state["active_section"] == "Similarity":
    st.header("Plagiarism & Similarity Detection")
    content = st.session_state["results"]
    if not content:
        st.info("No content loaded. Please use the Analyze tab first.")
    elif SentenceTransformer and cosine_similarity:
        st.markdown("Compare your text to another sample:")
        compare_text = st.text_area("Paste text to compare:")
        if st.button("Check Similarity") and compare_text:
            model = get_sentence_transformer()
            emb1 = model.encode([content])[0].reshape(1, -1)
            emb2 = model.encode([compare_text])[0].reshape(1, -1)
            sim = cosine_similarity(emb1, emb2)[0][0]
            st.success(f"Cosine Similarity: {sim:.2f}")
            if sim > 0.8:
                st.warning("High similarity detected! Possible plagiarism.")
    else:
        st.warning("sentence-transformers or scikit-learn not installed. Run: pip install sentence-transformers scikit-learn")

# --- Workflow Tab ---
if st.session_state["active_section"] == "Workflow":
    st.header("Customizable Workflows (Chain Tools)")
    st.info("Select a sequence of tools to apply. (Drag-and-drop coming soon!)")
    steps = st.multiselect("Choose steps:", ["OCR", "Translate", "Summarize", "NER", "TTS"])
    content = st.session_state["results"]
    workflow_result = content
    if st.button("Run Workflow") and content and steps:
        for step in steps:
            if step == "OCR" and Image and pytesseract:
                st.info("OCR already applied if image uploaded.")
            elif step == "Translate" and pipeline:
                translator = get_translator(lang)
                workflow_result = translator(workflow_result)[0]['translation_text']
            elif step == "Summarize" and pipeline:
                summarizer = get_summarizer()
                workflow_result = summarizer(workflow_result, max_length=130, min_length=30, do_sample=False)[0]["summary_text"]
            elif step == "NER" and nlp:
                doc = nlp(workflow_result)
                ents = [(ent.text, ent.label_) for ent in doc.ents]
                workflow_result += "\nEntities: " + ", ".join([f"{e[0]}({e[1]})" for e in ents])
            elif step == "TTS" and gTTS:
                tts = gTTS(workflow_result, lang=lang)
                tts_fp = BytesIO()
                tts.write_to_fp(tts_fp)
                tts_fp.seek(0)
                b64 = base64.b64encode(tts_fp.read()).decode()
                st.audio(f"data:audio/mp3;base64,{b64}", format="audio/mp3")
        st.success("Workflow result:")
        st.code(workflow_result, language="text")

# --- Timeline Extraction Tab ---
if st.session_state["active_section"] == "Timeline":
    st.header("Timeline Extraction & Visualization")
    content = st.session_state["results"]
    if not content or not nlp:
        st.info("No content loaded or spaCy not available.")
    else:
        doc = nlp(content)
        events = []
        for sent in doc.sents:
            if any(ent.label_ == "DATE" for ent in sent.ents):
                date = next((ent.text for ent in sent.ents if ent.label_ == "DATE"), None)
                events.append((date, sent.text))
        if events:
            st.markdown("**Timeline of Events:**")
            for date, event in sorted(events):
                st.markdown(f"- **{date}**: {event}")
        else:
            st.info("No dated events found in the text.")

# --- Topic Modeling Tab ---
if st.session_state["active_section"] == "Topics":
    st.header("Topic Modeling & Clustering")
    content = st.session_state["results"]
    if not content:
        st.info("No content loaded. Please use the Analyze tab first.")
    else:
        try:
            from sklearn.feature_extraction.text import CountVectorizer
            from sklearn.decomposition import LatentDirichletAllocation
            import numpy as np
            sentences = [s for s in content.split('.') if len(s.split()) > 3]
            if len(sentences) < 3:
                st.info("Not enough content for topic modeling.")
            else:
                vectorizer = CountVectorizer(stop_words='english')
                X = vectorizer.fit_transform(sentences)
                lda = LatentDirichletAllocation(n_components=3, random_state=42)
                lda.fit(X)
                words = np.array(vectorizer.get_feature_names_out())
                st.markdown("**Main Topics:**")
                for idx, topic in enumerate(lda.components_):
                    top_words = words[np.argsort(topic)[-5:]]
                    st.markdown(f"- Topic {idx+1}: {', '.join(top_words)}")
        except Exception as e:
            st.warning(f"Topic modeling requires scikit-learn and numpy: {e}")

# --- Explainability Tab ---
if st.session_state["active_section"] == "Explain":
    st.header("Explainability & Model Transparency")
    content = st.session_state["results"]
    ai_outputs = st.session_state["ai_outputs"]
    if not content or not ai_outputs:
        st.info("No analysis results to explain. Use the AI Tools tab first.")
    else:
        # Show model confidence for sentiment
        if "sentiment" in ai_outputs and "sentiment" in ai_outputs:
            st.markdown("**Sentiment Model Confidence:**")
            sentiment = get_sentiment()
            if sentiment:
                try:
                    sent = sentiment(content, return_all_scores=True)[0]
                    for s in sent:
                        st.write(f"{s['label']}: {s['score']:.2f}")
                except Exception as e:
                    st.warning(f"Could not get confidence scores: {e}")
        # Highlight influential words for NER
        if "ner" in ai_outputs and nlp:
            st.markdown("**Named Entities Highlighted:**")
            doc = nlp(content)
            highlighted = content
            for ent in doc.ents:
                highlighted = highlighted.replace(ent.text, f"**[{ent.text} ({ent.label_})]**")
            st.markdown(highlighted)
        # Show summary with important sentences bolded (rudimentary)
        if "summary" in ai_outputs:
            st.markdown("**Summary with Key Sentences:**")
            summary = ai_outputs["summary"]
            for sent in summary.split('. '):
                if any(word in sent for word in ["important", "key", "notable", "critical"]):
                    st.markdown(f"**{sent}**.")
                else:
                    st.markdown(f"{sent}.")

# --- Settings Tab ---
if st.session_state["active_section"] == "Settings":
    st.header("Settings & Privacy")
    st.markdown("- All processing is local unless you use a public API.\n- You can clear all data below.")
    if st.button("Clear All Data"):
        st.session_state.clear()
        st.success("All session data cleared.")

# --- History Tab ---
if st.session_state["active_section"] == "History":
    st.header("History")
    st.write("View your previous results and chat history.")
    if st.session_state.get("results"):
        st.markdown(f"**Last Analysis Result:**\n\n{st.session_state['results']}")
    if st.session_state.get("ai_outputs"):
        st.markdown("**AI Outputs:**")
        for k, v in st.session_state["ai_outputs"].items():
            st.markdown(f"- **{k.title()}:** {v}")
    if st.session_state.get("chat_history"):
        st.markdown("**Chat History:**")
        for sender, message in st.session_state["chat_history"]:
            st.markdown(f"**{sender}:** {message}")
    if st.session_state.get("batch_results"):
        st.markdown("**Batch Results:**")
        for fname, text in st.session_state["batch_results"]:
            st.markdown(f"**{fname}:**")
            st.code(text, language="text")
    else:
        st.write("No history yet.")

# --- Export Tab ---
if st.session_state["active_section"] == "Export":
    st.header("Export & Sharing")
    content = st.session_state["results"]
    ai_outputs = st.session_state["ai_outputs"]
    if not content:
        st.info("No content to export. Use the Analyze tab first.")
    else:
        st.markdown("**Export your results:**")
        # Export as Markdown
        if st.button("Export as Markdown"):
            md = f"# Analysis Result\n\n{content}\n\n## AI Outputs\n" + "\n".join([f"**{k.title()}**: {v}" for k, v in ai_outputs.items()])
            st.download_button("Download Markdown", md, file_name="obsidian_output.md")
        # Export as PDF (requires fpdf)
        try:
            from fpdf import FPDF
            if st.button("Export as PDF"):
                pdf = FPDF()
                pdf.add_page()
                pdf.set_font("Arial", size=12)
                pdf.multi_cell(0, 10, content)
                for k, v in ai_outputs.items():
                    pdf.multi_cell(0, 10, f"{k.title()}: {v}")
                pdf_fp = BytesIO()
                pdf.output(pdf_fp)
                pdf_fp.seek(0)
                st.download_button("Download PDF", pdf_fp, file_name="obsidian_output.pdf")
        except Exception as e:
            st.info("Install fpdf for PDF export: pip install fpdf")
        # Export as Word (requires python-docx)
        try:
            from docx import Document
            if st.button("Export as Word"):
                doc = Document()
                doc.add_heading("Analysis Result", 0)
                doc.add_paragraph(content)
                doc.add_heading("AI Outputs", level=1)
                for k, v in ai_outputs.items():
                    doc.add_paragraph(f"{k.title()}: {v}")
                doc_fp = BytesIO()
                doc.save(doc_fp)
                doc_fp.seek(0)
                st.download_button("Download Word", doc_fp, file_name="obsidian_output.docx")
        except Exception as e:
            st.info("Install python-docx for Word export: pip install python-docx")

# --- Accessibility Tab ---
if st.session_state["active_section"] == "Accessibility":
    st.header("Accessibility & Display Settings")
    st.markdown("**Adjust font size and enable dyslexia-friendly font.**")
    font_size = st.slider("Font Size", 12, 32, 16)
    dyslexia_font = st.checkbox("Enable Dyslexia-Friendly Font")
    st.markdown(f"<style>body, .stApp, .stMarkdown, .stTextInput, .stTextArea, .stButton, .stSidebar {{ font-size: {font_size}px !important; {'font-family: OpenDyslexic, Arial, sans-serif !important;' if dyslexia_font else ''} }}</style>", unsafe_allow_html=True)
    st.info("Font size and font will update on next rerun or tab switch.")

# --- User Profiles & Saved Sessions Tab ---
if st.session_state["active_section"] == "User Profiles":
    st.header("User Profiles & Saved Sessions")
    username = st.text_input("Enter your username to save/load your session:")
    if st.button("Save Session") and username:
        st.session_state[f"profile_{username}"] = {
            "results": st.session_state["results"],
            "ai_outputs": st.session_state["ai_outputs"],
            "chat_history": st.session_state["chat_history"]
        }
        st.success(f"Session saved for {username}.")
    if st.button("Load Session") and username and f"profile_{username}" in st.session_state:
        profile = st.session_state[f"profile_{username}"]
        st.session_state["results"] = profile["results"]
        st.session_state["ai_outputs"] = profile["ai_outputs"]
        st.session_state["chat_history"] = profile["chat_history"]
        st.success(f"Session loaded for {username}.")
    st.info("Sessions are saved locally in your browser session.")

# --- Real-Time Collaboration Placeholder Tab ---
if st.session_state["active_section"] == "Real-Time Collaboration":
    st.header("Real-Time Collaboration (Coming Soon)")
    st.info("This feature will allow multiple users to analyze and discuss the same document in real time.")
    st.markdown("If you want to help build this, open an issue or PR!")

# --- Data Visualizations Tab ---
if st.session_state["active_section"] == "Data Visualizations":
    st.header("Data Visualizations")
    import matplotlib.pyplot as plt
    import collections
    content = st.session_state["results"]
    ai_outputs = st.session_state["ai_outputs"]
    if not content or not ai_outputs:
        st.info("No analysis results to visualize. Use the AI Tools tab first.")
    else:
        # Entity frequency
        if "ner" in ai_outputs and ai_outputs["ner"]:
            st.subheader("Entity Frequency")
            ents = [ent for ent, label in ai_outputs["ner"]]
            counter = collections.Counter(ents)
            fig, ax = plt.subplots()
            ax.bar(counter.keys(), counter.values(), color="#001F3F")
            plt.xticks(rotation=45)
            st.pyplot(fig)
        # Sentiment pie chart
        if "sentiment" in ai_outputs and "POSITIVE" in ai_outputs["sentiment"] or "NEGATIVE" in ai_outputs["sentiment"]:
            st.subheader("Sentiment Distribution")
            labels = ["POSITIVE", "NEGATIVE", "NEUTRAL"]
            values = [ai_outputs["sentiment"].count(l) for l in labels]
            fig, ax = plt.subplots()
            ax.pie(values, labels=labels, autopct='%1.1f%%', colors=["#0074D9", "#FF4136", "#AAAAAA"])
            st.pyplot(fig)
        # Topic bar chart
        if "topics" in ai_outputs and ai_outputs["topics"]:
            st.subheader("Topic Prevalence")
            topics = ai_outputs["topics"]
            fig, ax = plt.subplots()
            ax.bar([f"Topic {i+1}" for i in range(len(topics))], [len(t) for t in topics], color="#39CCCC")
            st.pyplot(fig)

# --- Theme Picker Tab ---
if st.session_state["active_section"] == "Theme Picker":
    st.header("Customizable Themes")
    theme = st.selectbox("Choose a theme:", ["Black & Navy Blue", "Classic Dark", "Light", "Solarized"])
    if theme == "Black & Navy Blue":
        st.markdown('<style>body, .stApp { background-color: #000014; color: #E4E6EB; } .stTabs [aria-selected="true"] { background: #001F3F; color: #FFD700; }</style>', unsafe_allow_html=True)
    elif theme == "Classic Dark":
        st.markdown('<style>body, .stApp { background-color: #18191A; color: #E4E6EB; } .stTabs [aria-selected="true"] { background: #3A3B3C; color: #FFD700; }</style>', unsafe_allow_html=True)
    elif theme == "Light":
        st.markdown('<style>body, .stApp { background-color: #FFFFFF; color: #222; } .stTabs [aria-selected="true"] { background: #E0E0E0; color: #0074D9; }</style>', unsafe_allow_html=True)
    elif theme == "Solarized":
        st.markdown('<style>body, .stApp { background-color: #002B36; color: #839496; } .stTabs [aria-selected="true"] { background: #073642; color: #B58900; }</style>', unsafe_allow_html=True)
    st.info("Theme will update on next rerun or tab switch.")

# --- EVOLUTION/JOURNEY TIMELINE SECTION ---
if "timeline_memories" not in st.session_state:
    st.session_state["timeline_memories"] = []

if st.session_state["active_section"] == "Journey":
    with glass_card():
        st.markdown('''
            <style>
            .timeline {
                display: flex;
                flex-direction: column;
                gap: 2.5rem;
                margin-top: 1.5rem;
                margin-bottom: 1.5rem;
                position: relative;
            }
            .timeline-era {
                display: flex;
                align-items: flex-start;
                gap: 2.2rem;
                background: rgba(255,255,255,0.18);
                border-radius: 22px;
                box-shadow: 0 2px 16px #0057B822;
                padding: 1.5rem 2rem;
                animation: fadein 0.7s cubic-bezier(.4,2,.6,1);
                position: relative;
                border-left: 6px solid #FFD700;
            }
            .timeline-era:hover {
                background: rgba(0,87,184,0.18);
                box-shadow: 0 4px 32px #FFD70044;
                transform: scale(1.02);
            }
            .timeline-mockup {
                min-width: 80px; max-width: 80px;
                min-height: 80px; max-height: 80px;
                display: flex; align-items: center; justify-content: center;
                font-size: 2.8rem;
                background: rgba(0,87,184,0.13);
                border-radius: 18px;
                box-shadow: 0 2px 8px #FFD70022;
                margin-right: 1.2rem;
            }
            .timeline-content {
                flex: 1;
            }
            .timeline-title {
                font-size: 1.25rem;
                font-weight: 700;
                color: #FFD700;
                margin-bottom: 0.3rem;
            }
            .timeline-date {
                font-size: 1.01rem;
                color: #b3d4fc;
                margin-bottom: 0.7rem;
            }
            .timeline-features {
                font-size: 1.05rem;
                color: #0057B8;
                margin-bottom: 0.5rem;
            }
            .timeline-badge {
                display: inline-block;
                background: linear-gradient(90deg, #FFD700 0%, #0057B8 100%);
                color: #232946;
                font-size: 0.98rem;
                font-weight: 600;
                border-radius: 8px;
                padding: 0.2rem 0.8rem;
                margin-right: 0.5rem;
                margin-bottom: 0.2rem;
                box-shadow: 0 2px 8px #FFD70022;
            }
            .timeline-memory {
                background: rgba(255,255,255,0.13);
                border-radius: 12px;
                padding: 0.7rem 1.2rem;
                color: #0057B8;
                font-size: 1.01rem;
                margin-top: 0.5rem;
                margin-bottom: 0.5rem;
                box-shadow: 0 2px 8px #FFD70022;
            }
            </style>
        ''', unsafe_allow_html=True)
        st.markdown('<div class="timeline">', unsafe_allow_html=True)
        # Timeline eras (mockups, descriptions, features)
        timeline_eras = [
            {
                "title": "v1.0 – Classic Dashboard",
                "date": "Spring 2024",
                "mockup": "🗂️",
                "desc": "The original Obsidian Protocol: simple, white, with tabs for each feature.",
                "features": ["Text/PDF/URL input", "Summary", "Sentiment", "NER"],
                "badges": ["First Release"]
            },
            {
                "title": "v2.0 – Glassmorphic Revolution",
                "date": "Summer 2024",
                "mockup": "🪟",
                "desc": "Major redesign: glass cards, blue/gold palette, micro-interactions.",
                "features": ["Image/audio input", "TTS", "Translation", "Batch", "Knowledge Graph"],
                "badges": ["Glassmorphic", "Luxury"]
            },
            {
                "title": "v3.0 – ChatGPT Mode",
                "date": "Summer 2024",
                "mockup": "💬",
                "desc": "Conversational, ChatGPT-style interface. Drag-and-drop, voice, AI suggestions.",
                "features": ["Central chat", "Live preview", "Animated suggestions", "Confetti"],
                "badges": ["Conversational", "Next-Gen"]
            },
            {
                "title": "v4.0 – The Living App",
                "date": "Future",
                "mockup": "🧬",
                "desc": "Live collaboration, evolution timeline, AI personas, and more.",
                "features": ["Collaboration", "Personas", "Achievements", "Animated timeline"],
                "badges": ["Living App", "Legendary"]
            }
        ]
        for era in timeline_eras:
            st.markdown(f'''<div class="timeline-era">
                <div class="timeline-mockup">{era['mockup']}</div>
                <div class="timeline-content">
                    <div class="timeline-title">{era['title']}</div>
                    <div class="timeline-date">{era['date']}</div>
                    <div class="timeline-features">{era['desc']}</div>
                    {''.join([f'<span class="timeline-badge">{b}</span>' for b in era['badges']])}
                    <ul style="margin:0.5rem 0 0.5rem 1.2rem; color:#0057B8;">
                        {''.join([f'<li>{f}</li>' for f in era['features']])}
                    </ul>
                    {''.join([f'<div class="timeline-memory">📝 {m}</div>' for m in st.session_state["timeline_memories"] if m.get("era") == era["title"]])}
                </div>
            </div>''', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        # Add user memory
        st.markdown("**Add your own memory or comment to the timeline:**")
        memory_text = st.text_input("Share a memory, feedback, or milestone:")
        selected_era = st.selectbox("Which era does this memory belong to?", [era["title"] for era in timeline_eras])
        if st.button("Add to Timeline") and memory_text:
            st.session_state["timeline_memories"].append({"era": selected_era, "text": memory_text})
            st.success("Memory added! Refresh to see it on the timeline.")

# Footer
st.markdown("<hr style='margin-top:2em;margin-bottom:1em'>", unsafe_allow_html=True)
st.caption("Obsidian Protocol v2.0 | World-Changing Streamlit UI | All media types & AI tools | Free forever | by AI")

st.markdown('</div>', unsafe_allow_html=True)

# --- LEGACY VAULT & AWE-INSPIRING CENTERPIECE ---
if ("Legacy Vault", "🏛️") not in SECTIONS:
    SECTIONS.append(("Legacy Vault", "🏛️"))

if "hall_of_fame" not in st.session_state:
    st.session_state["hall_of_fame"] = []

# Cinematic intro animation for first-time users (DISABLED to prevent blocking)
if "cinematic_intro" not in st.session_state:
    st.session_state["cinematic_intro"] = True  # Skip intro, set to True by default

# Legacy Vault section
if st.session_state["active_section"] == "Legacy Vault":
    with glass_card():
        st.markdown('''
            <style>
            .legacy-timeline {
                display: flex; flex-direction: column; gap: 2.5rem; margin-top: 1.5rem; margin-bottom: 1.5rem; position: relative;
            }
            .legacy-era {
                display: flex; align-items: flex-start; gap: 2.2rem;
                background: rgba(255,255,255,0.18); border-radius: 22px;
                box-shadow: 0 2px 16px #0057B822; padding: 1.5rem 2rem;
                animation: fadein 0.7s cubic-bezier(.4,2,.6,1);
                position: relative; border-left: 6px solid #FFD700;
                transition: background 0.3s, box-shadow 0.3s, transform 0.3s;
            }
            .legacy-era:hover {
                background: rgba(0,87,184,0.18); box-shadow: 0 4px 32px #FFD70044; transform: scale(1.02);
            }
            .legacy-mockup {
                min-width: 100px; max-width: 100px; min-height: 100px; max-height: 100px;
                display: flex; align-items: center; justify-content: center;
                font-size: 3.2rem; background: rgba(0,87,184,0.13);
                border-radius: 22px; box-shadow: 0 2px 8px #FFD70022; margin-right: 1.2rem;
                position: relative; overflow: hidden;
            }
            .legacy-mockup .parallax {
                position: absolute; width: 100%; height: 100%; pointer-events: none;
                background: radial-gradient(circle at 30% 40%, #FFD70033 0%, transparent 80%),
                            radial-gradient(circle at 70% 60%, #0057B822 0%, transparent 80%);
                animation: parallaxmove 4s infinite alternate;
            }
            @keyframes parallaxmove {
                0% { background-position: 0% 0%, 100% 100%; }
                100% { background-position: 100% 100%, 0% 0%; }
            }
            .legacy-content { flex: 1; }
            .legacy-title { font-size: 1.35rem; font-weight: 700; color: #FFD700; margin-bottom: 0.3rem; }
            .legacy-date { font-size: 1.01rem; color: #b3d4fc; margin-bottom: 0.7rem; }
            .legacy-features { font-size: 1.05rem; color: #0057B8; margin-bottom: 0.5rem; }
            .legacy-badge { display: inline-block; background: linear-gradient(90deg, #FFD700 0%, #0057B8 100%); color: #232946; font-size: 0.98rem; font-weight: 600; border-radius: 8px; padding: 0.2rem 0.8rem; margin-right: 0.5rem; margin-bottom: 0.2rem; box-shadow: 0 2px 8px #FFD70022; }
            .legacy-signature { background: rgba(255,255,255,0.13); border-radius: 12px; padding: 0.7rem 1.2rem; color: #0057B8; font-size: 1.01rem; margin-top: 0.5rem; margin-bottom: 0.5rem; box-shadow: 0 2px 8px #FFD70022; }
            .legacy-hof { margin-top: 2.5rem; background: rgba(0,87,184,0.13); border-radius: 18px; padding: 1.2rem 2rem; box-shadow: 0 2px 8px #FFD70022; }
            </style>
        ''', unsafe_allow_html=True)
        st.markdown('<div class="legacy-timeline">', unsafe_allow_html=True)
        legacy_eras = [
            {"title": "v1.0 – Classic Dashboard", "date": "Spring 2024", "mockup": "🗂️", "desc": "The original: simple, white, with tabs.", "features": ["Text/PDF/URL input", "Summary", "Sentiment", "NER"], "badges": ["First Release"]},
            {"title": "v2.0 – Glassmorphic Revolution", "date": "Summer 2024", "mockup": "🪟", "desc": "Major redesign: glass cards, blue/gold palette, micro-interactions.", "features": ["Image/audio input", "TTS", "Translation", "Batch", "Knowledge Graph"], "badges": ["Glassmorphic", "Luxury"]},
            {"title": "v3.0 – ChatGPT Mode", "date": "Summer 2024", "mockup": "💬", "desc": "Conversational, ChatGPT-style interface. Drag-and-drop, voice, AI suggestions.", "features": ["Central chat", "Live preview", "Animated suggestions", "Confetti"], "badges": ["Conversational", "Next-Gen"]},
            {"title": "v4.0 – The Living App", "date": "Future", "mockup": "🧬", "desc": "Live collaboration, evolution timeline, AI personas, and more.", "features": ["Collaboration", "Personas", "Achievements", "Animated timeline"], "badges": ["Living App", "Legendary"]}
        ]
        for era in legacy_eras:
            st.markdown(f'''<div class="legacy-era">
                <div class="legacy-mockup">{era['mockup']}<div class="parallax"></div></div>
                <div class="legacy-content">
                    <div class="legacy-title">{era['title']}</div>
                    <div class="legacy-date">{era['date']}</div>
                    <div class="legacy-features">{era['desc']}</div>
                    {''.join([f'<span class="legacy-badge">{b}</span>' for b in era['badges']])}
                    <ul style="margin:0.5rem 0 0.5rem 1.2rem; color:#0057B8;">
                        {''.join([f'<li>{f}</li>' for f in era['features']])}
                    </ul>
                    {''.join([f'<div class="legacy-signature">📝 {s}</div>' for s in st.session_state["hall_of_fame"] if s.get("era") == era["title"]])}
                </div>
            </div>''', unsafe_allow_html=True)
        st.markdown('</div>', unsafe_allow_html=True)
        # User signature
        st.markdown("**Sign the Vault:**")
        signature_text = st.text_input("Leave your name or message for the Hall of Fame:")
        selected_era = st.selectbox("Which era do you want to sign?", [era["title"] for era in legacy_eras])
        if st.button("Sign Vault") and signature_text:
            st.session_state["hall_of_fame"].append({"era": selected_era, "text": signature_text})
            st.success("Signature added! Refresh to see it in the Vault.")

# --- UPGRADED CENTERPIECE: 3D GLASS, PARALLAX, LIVING BACKGROUND ---
st.markdown('''
    <style>
    .centerpiece-glass {
        margin-left: 90px; margin-top: 2.5rem; margin-bottom: 2.5rem;
        padding: 3.5rem 3rem 3rem 3rem;
        background: linear-gradient(120deg, #1a1f2b 0%, #232946 50%, #2e3a59 100%);
        border-radius: 48px;
        box-shadow: 0 16px 64px 0 #0057B844, 0 2px 16px #FFD70033;
        border: 3px solid rgba(255,255,255,0.22);
        min-height: 420px; max-width: 900px;
        margin-right: auto;
        position: relative;
        overflow: hidden;
        animation: fadein 1.2s cubic-bezier(.4,2,.6,1);
        perspective: 1200px;
        transition: box-shadow 0.22s, background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1);
    }
    .centerpiece-glass .glass-inner {
        background: rgba(255,255,255,0.18);
        border-radius: 40px;
        box-shadow: 0 8px 32px 0 #0057B822, 0 2px 8px #FFD70022;
        padding: 2.5rem 2rem 2rem 2rem;
        min-height: 320px;
        transition: box-shadow 0.22s cubic-bezier(.4,2,.6,1), background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1);
        will-change: box-shadow, background, transform, opacity;
        position: relative;
        animation: glassbreath 2.5s infinite alternate;
    }
    @keyframes glassbreath {
        0% { box-shadow: 0 8px 32px 0 #0057B822, 0 2px 8px #FFD70022; transform: scale(1) rotateY(0deg); }
        100% { box-shadow: 0 16px 64px 0 #FFD70044, 0 4px 16px #0057B844; transform: scale(1.015) rotateY(2deg); }
    }
    .centerpiece-glass .living-bg {
        position: absolute; top:0; left:0; width:100%; height:100%; z-index:0;
        pointer-events: none;
        background: radial-gradient(circle at 20% 40%, #00c3ff33 0%, #0057b822 60%, transparent 100%),
                    radial-gradient(circle at 80% 60%, #ffd70033 0%, #ffb30022 60%, transparent 100%);
        animation: livingbgmove 12s ease-in-out infinite alternate;
        opacity: 0.7;
    }
    @keyframes livingbgmove {
        0% { background-position: 0% 0%, 100% 100%; }
        100% { background-position: 100% 100%, 0% 0%; }
    }
    .centerpiece-glass .light-reflection {
        position: absolute; top: 0; left: 0; width: 100%; height: 100%;
        background: linear-gradient(120deg, #fff3 0%, #fff0 100%);
        opacity: 0.18;
        pointer-events: none;
        mix-blend-mode: screen;
        animation: reflectionmove 6s infinite alternate;
    }
    @keyframes reflectionmove {
        0% { opacity: 0.12; }
        100% { opacity: 0.22; }
    }
    </style>
''', unsafe_allow_html=True)

# Only show centerpiece on main/chat section (not Journey/Legacy Vault)
if st.session_state["active_section"] not in ["Journey", "Legacy Vault"]:
    st.markdown('<div class="centerpiece-glass" id="centerpiece">', unsafe_allow_html=True)
    st.markdown('<div class="living-bg"></div>', unsafe_allow_html=True)
    st.markdown('<div class="light-reflection"></div>', unsafe_allow_html=True)
    st.markdown('<div class="glass-inner">', unsafe_allow_html=True)
    # (Insert your chat/search input and results here)
    st.markdown('</div>', unsafe_allow_html=True)
    st.markdown('</div>', unsafe_allow_html=True)

# --- EMOTIONAL, JOYFUL, 3D TYPING EXPERIENCE ---
st.markdown('''
    <style>
    .centerpiece-glass {
        margin-left: 90px; margin-top: 2.5rem; margin-bottom: 2.5rem;
        padding: 3.5rem 3rem 3rem 3rem;
        background: linear-gradient(120deg, #1a1f2b 0%, #232946 50%, #2e3a59 100%);
        border-radius: 48px;
        box-shadow: 0 16px 64px 0 #0057B844, 0 2px 16px #FFD70033;
        border: 3px solid rgba(255,255,255,0.22);
        min-height: 420px; max-width: 900px;
        margin-right: auto;
        position: relative;
        overflow: hidden;
        animation: fadein 1.2s cubic-bezier(.4,2,.6,1);
        perspective: 1200px;
        transition: box-shadow 0.22s, background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1);
    }
    .centerpiece-glass .glass-inner {
        background: rgba(255,255,255,0.18);
        border-radius: 40px;
        box-shadow: 0 8px 32px 0 #0057B822, 0 2px 8px #FFD70022;
        padding: 2.5rem 2rem 2rem 2rem;
        min-height: 320px;
        transition: box-shadow 0.22s cubic-bezier(.4,2,.6,1), background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1);
        will-change: box-shadow, background, transform, opacity;
        position: relative;
        animation: glassbreath 2.5s infinite alternate;
    }
    @keyframes glassbreath {
        0% { box-shadow: 0 8px 32px 0 #0057B822, 0 2px 8px #FFD70022; transform: scale(1) rotateY(0deg); }
        100% { box-shadow: 0 16px 64px 0 #FFD70044, 0 4px 16px #0057B844; transform: scale(1.015) rotateY(2deg); }
    }
    .centerpiece-glass .living-bg {
        position: absolute; top:0; left:0; width:100%; height:100%; z-index:0;
        pointer-events: none;
        background: radial-gradient(circle at 20% 40%, #00c3ff33 0%, #0057b822 60%, transparent 100%),
                    radial-gradient(circle at 80% 60%, #ffd70033 0%, #ffb30022 60%, transparent 100%);
        animation: livingbgmove 12s ease-in-out infinite alternate;
        opacity: 0.7;
    }
    @keyframes livingbgmove {
        0% { background-position: 0% 0%, 100% 100%; }
        100% { background-position: 100% 100%, 0% 0%; }
    }
    .centerpiece-glass .light-reflection {
        position: absolute; top: 0; left: 0; width: 100%; height: 100%;
        background: linear-gradient(120deg, #fff3 0%, #fff0 100%);
        opacity: 0.18;
        pointer-events: none;
        mix-blend-mode: screen;
        animation: reflectionmove 6s infinite alternate;
    }
    @keyframes reflectionmove {
        0% { opacity: 0.12; }
        100% { opacity: 0.22; }
    }
    .chat-input {
        flex: 1;
        font-size: 1.18rem;
        padding: 1.1rem 1.5rem;
        border-radius: 18px;
        border: 2.5px solid #FFD700;
        background: rgba(255,255,255,0.7);
        color: #232946;
        box-shadow: 0 2px 16px #FFD70033;
        outline: none;
        transition: border 0.22s cubic-bezier(.4,2,.6,1), box-shadow 0.22s cubic-bezier(.4,2,.6,1), background 0.22s cubic-bezier(.4,2,.6,1), transform 0.18s cubic-bezier(.4,2,.6,1);
        will-change: border, box-shadow, background, transform;
        position: relative;
        z-index: 2;
    }
    .chat-input:focus {
        border: 2.5px solid #0057B8;
        box-shadow: 0 0 0 8px #0057B822;
        background: rgba(255,255,255,0.9);
        animation: inputpulse 0.7s;
    }
    @keyframes inputpulse {
        0% { box-shadow: 0 0 0 0 #FFD70044; }
        100% { box-shadow: 0 0 0 8px #0057B822; }
    }
    .chat-input.typing {
        animation: typingpulse 0.18s;
    }
    @keyframes typingpulse {
        0% { background: rgba(255,255,255,0.7); }
        100% { background: rgba(255,255,255,0.9); }
    }
    .chat-input-sparkle {
        position: absolute; right: 18px; top: 50%; transform: translateY(-50%);
        font-size: 1.3rem; color: #FFD700; opacity: 0.7;
        pointer-events: none;
        animation: sparklepop 0.5s;
    }
    @keyframes sparklepop {
        0% { opacity: 0; transform: scale(0.7) translateY(-50%); }
        50% { opacity: 1; transform: scale(1.2) translateY(-50%); }
        100% { opacity: 0; transform: scale(0.7) translateY(-50%); }
    }
    .chat-send-btn:active {
        animation: sendwave 0.4s;
    }
    @keyframes sendwave {
        0% { box-shadow: 0 2px 16px #FFD70077; }
        50% { box-shadow: 0 8px 32px #FFD70077; }
        100% { box-shadow: 0 2px 16px #FFD70077; }
    }
    .positive-feedback {
        font-size: 1.12rem; color: #FFD700; margin-top: 0.7rem; text-align: center;
        animation: fadein 0.7s cubic-bezier(.4,2,.6,1);
        font-weight: 600;
        letter-spacing: 0.01em;
        text-shadow: 0 2px 8px #0057B822;
    }
    </style>
''', unsafe_allow_html=True)

# (JS for 3D parallax and input sparkle would be added in a full web app; Streamlit is limited, but CSS/HTML is in place)

# Example: Show positive feedback after send
if "show_positive_feedback" not in st.session_state:
    st.session_state["show_positive_feedback"] = False
if st.session_state["show_positive_feedback"]:
    st.markdown('<div class="positive-feedback">Great thought! ✨</div>', unsafe_allow_html=True)
    st.session_state["show_positive_feedback"] = False

# --- ULTIMATE EMOTIONAL, 3D, INTERACTIVE TYPING EXPERIENCE ---

# Add toggle in settings for effects
if "effects_enabled" not in st.session_state:
    st.session_state["effects_enabled"] = True
if st.session_state["active_section"] == "Settings":
    st.markdown("**Personalization & Effects**")
    st.session_state["effects_enabled"] = st.checkbox("Enable 3D, emoji, and sound effects (recommended)", value=st.session_state["effects_enabled"])

# Only apply effects if enabled
if st.session_state.get("effects_enabled", True):
    st.markdown('''
        <style>
        .centerpiece-glass {
            transition: box-shadow 0.22s, background 0.5s, transform 0.22s cubic-bezier(.4,2,.6,1);
        }
        </style>
        <script>
        // 3D parallax/tilt effect
        const card = window.parent.document.getElementById('centerpiece');
        if(card){
            card.onmousemove = function(e){
                const rect = card.getBoundingClientRect();
                const x = e.clientX - rect.left;
                const y = e.clientY - rect.top;
                const cx = rect.width/2, cy = rect.height/2;
                const dx = (x-cx)/cx, dy = (y-cy)/cy;
                card.style.transform = `rotateY(${-dx*8}deg) rotateX(${dy*8}deg) scale(1.01)`;
            };
            card.onmouseleave = function(){
                card.style.transform = '';
            };
        }
        // Emoji confetti on typing/send
        window.emojiConfetti = function(){
            const input = window.parent.document.querySelector('.chat-input');
            if(!input) return;
            for(let i=0;i<6;i++){
                const emoji = document.createElement('span');
                emoji.innerText = ['✨','🎉','😊','🥳','💙','⭐'][Math.floor(Math.random()*6)];
                emoji.style.position = 'absolute';
                emoji.style.left = (50+Math.random()*30-15)+'%';
                emoji.style.top = (60+Math.random()*10-5)+'%';
                emoji.style.fontSize = (1.2+Math.random()*0.8)+'rem';
                emoji.style.opacity = 1;
                emoji.style.transition = 'all 1.2s cubic-bezier(.4,2,.6,1)';
                emoji.style.pointerEvents = 'none';
                input.parentElement.appendChild(emoji);
                setTimeout(()=>{
                    emoji.style.transform = `translateY(-60px) scale(${1.2+Math.random()*0.5})`;
                    emoji.style.opacity = 0;
                }, 10);
                setTimeout(()=>{emoji.remove();}, 1200);
            }
        };
        // Sound cue on send
        window.playSendSound = function(){
            const ctx = new (window.AudioContext||window.webkitAudioContext)();
            const o = ctx.createOscillator();
            const g = ctx.createGain();
            o.type = 'triangle';
            o.frequency.value = 660;
            g.gain.value = 0.08;
            o.connect(g); g.connect(ctx.destination);
            o.start();
            o.frequency.linearRampToValueAtTime(880, ctx.currentTime+0.18);
            g.gain.linearRampToValueAtTime(0, ctx.currentTime+0.22);
            o.stop(ctx.currentTime+0.22);
        };
        </script>
    ''', unsafe_allow_html=True)

# Example: On send, trigger emoji confetti, sound, and positive feedback
if "show_positive_feedback" not in st.session_state:
    st.session_state["show_positive_feedback"] = False
if st.session_state["show_positive_feedback"]:
    st.markdown('<div class="positive-feedback">Great thought! ✨</div>', unsafe_allow_html=True)
    st.session_state["show_positive_feedback"] = False
    st.markdown('''<script>window.emojiConfetti();window.playSendSound();</script>''', unsafe_allow_html=True)

# --- RARE SUPER REACTIONS, ANIMATED BACKGROUNDS, MORE SOUNDS ---

if "message_count" not in st.session_state:
    st.session_state["message_count"] = 0

# Add more sound options and rare super reactions
if st.session_state.get("effects_enabled", True):
    st.markdown('''
        <script>
        // Golden confetti super reaction
        window.superConfetti = function(){
            const input = window.parent.document.querySelector('.chat-input');
            if(!input) return;
            for(let i=0;i<18;i++){
                const emoji = document.createElement('span');
                emoji.innerText = ['✨','🎉','⭐','💛','🥇','🏆'][Math.floor(Math.random()*6)];
                emoji.style.position = 'absolute';
                emoji.style.left = (30+Math.random()*40)+'%';
                emoji.style.top = (60+Math.random()*10-5)+'%';
                emoji.style.fontSize = (1.5+Math.random()*1.2)+'rem';
                emoji.style.opacity = 1;
                emoji.style.transition = 'all 1.8s cubic-bezier(.4,2,.6,1)';
                emoji.style.pointerEvents = 'none';
                input.parentElement.appendChild(emoji);
                setTimeout(()=>{
                    emoji.style.transform = `translateY(-120px) scale(${1.5+Math.random()*0.7}) rotate(${Math.random()*360}deg)`;
                    emoji.style.opacity = 0;
                }, 10);
                setTimeout(()=>{emoji.remove();}, 1800);
            }
        };
        // Gentle chime for positive feedback
        window.playChime = function(){
            const ctx = new (window.AudioContext||window.webkitAudioContext)();
            const o = ctx.createOscillator();
            const g = ctx.createGain();
            o.type = 'sine';
            o.frequency.value = 880;
            g.gain.value = 0.09;
            o.connect(g); g.connect(ctx.destination);
            o.start();
            o.frequency.linearRampToValueAtTime(1320, ctx.currentTime+0.22);
            g.gain.linearRampToValueAtTime(0, ctx.currentTime+0.32);
            o.stop(ctx.currentTime+0.32);
        };
        // Rare celebration sound
        window.playCelebration = function(){
            const ctx = new (window.AudioContext||window.webkitAudioContext)();
            const o1 = ctx.createOscillator();
            const o2 = ctx.createOscillator();
            const g = ctx.createGain();
            o1.type = 'triangle'; o2.type = 'sine';
            o1.frequency.value = 660; o2.frequency.value = 990;
            g.gain.value = 0.12;
            o1.connect(g); o2.connect(g); g.connect(ctx.destination);
            o1.start(); o2.start();
            o1.frequency.linearRampToValueAtTime(1320, ctx.currentTime+0.32);
            o2.frequency.linearRampToValueAtTime(1760, ctx.currentTime+0.32);
            g.gain.linearRampToValueAtTime(0, ctx.currentTime+0.38);
            o1.stop(ctx.currentTime+0.38); o2.stop(ctx.currentTime+0.38);
        };
        // Animated background pulse
        window.bgPulse = function(){
            const bg = window.parent.document.querySelector('.living-bg');
            if(bg){
                bg.style.transition = 'filter 0.7s cubic-bezier(.4,2,.6,1)';
                bg.style.filter = 'brightness(1.3) blur(8px)';
                setTimeout(()=>{bg.style.filter='';}, 700);
            }
        };
        </script>
    ''', unsafe_allow_html=True)

# Example: On send, trigger super reaction every 10th message
if "show_positive_feedback" not in st.session_state:
    st.session_state["show_positive_feedback"] = False
if st.session_state["show_positive_feedback"]:
    st.session_state["message_count"] += 1
    if st.session_state["message_count"] % 10 == 0:
        st.markdown('<div class="positive-feedback">Legendary! You hit a milestone! 🏆</div>', unsafe_allow_html=True)
        st.markdown('''<script>window.superConfetti();window.playCelebration();window.bgPulse();</script>''', unsafe_allow_html=True)
    else:
        st.markdown('<div class="positive-feedback">Great thought! ✨</div>', unsafe_allow_html=True)
        st.markdown('''<script>window.emojiConfetti();window.playChime();window.bgPulse();</script>''', unsafe_allow_html=True)
    st.session_state["show_positive_feedback"] = False

# --- ENHANCED ANIMATIONS: CHAT, BUTTONS, SIDEBAR, BACKGROUND ---
st.markdown('''
    <style>
    /* Animate chat bubbles */
    .chat-bubble {
        animation: chatbubblein 0.7s cubic-bezier(.4,2,.6,1);
        transition: background 0.18s, box-shadow 0.18s, transform 0.18s;
    }
    @keyframes chatbubblein {
        0% { opacity: 0; transform: translateY(24px) scale(0.96); }
        60% { opacity: 1; transform: translateY(-6px) scale(1.04); }
        100% { opacity: 1; transform: none; }
    }
    .chat-bubble.user {
        animation-delay: 0.1s;
    }
    .chat-bubble.ai {
        animation-delay: 0.2s;
    }
    /* Animate send button */
    .chat-send-btn {
        animation: sendpulsein 0.7s cubic-bezier(.4,2,.6,1);
    }
    @keyframes sendpulsein {
        0% { opacity: 0; transform: scale(0.8); }
        100% { opacity: 1; transform: scale(1); }
    }
    .chat-send-btn:hover {
        box-shadow: 0 0 16px #FFD70088, 0 2px 16px #0057B855;
        animation: sendpulse 1.2s infinite alternate;
    }
    @keyframes sendpulse {
        0% { box-shadow: 0 0 16px #FFD70088, 0 2px 16px #0057B855; }
        100% { box-shadow: 0 0 32px #FFD700cc, 0 4px 32px #0057B8cc; }
    }
    .chat-send-btn:active {
        animation: sendbounce 0.3s;
    }
    @keyframes sendbounce {
        0% { transform: scale(1); }
        40% { transform: scale(1.12) translateY(-4px); }
        100% { transform: scale(1); }
    }
    /* Sidebar avatar/profile pulse and glow */
    .sidebar-avatar {
        transition: box-shadow 0.3s, border 0.3s;
    }
    .sidebar-avatar:hover {
        box-shadow: 0 0 0 12px #FFD70044, 0 2px 12px #FFD70033;
        border: 2.5px solid #FFF;
        animation: avatarpulse 1.5s infinite alternate;
    }
    @keyframes avatarpulse {
        0% { box-shadow: 0 0 0 8px #FFD70033, 0 2px 12px #FFD70033; }
        100% { box-shadow: 0 0 0 16px #FFD70066, 0 4px 24px #FFD70066; }
    }
    /* Floating animated sparkles/particles in background */
    .magic-sparkle {
        position: fixed; pointer-events: none; z-index: 1;
        width: 100vw; height: 100vh; left: 0; top: 0;
        overflow: hidden;
    }
    .magic-sparkle span {
        position: absolute;
        border-radius: 50%;
        opacity: 0.7;
        pointer-events: none;
        animation: sparklefloat 6s linear infinite;
        background: linear-gradient(135deg, #FFD700 0%, #0057B8 100%);
    }
    @keyframes sparklefloat {
        0% { transform: translateY(0) scale(0.7); opacity: 0.7; }
        60% { opacity: 1; }
        100% { transform: translateY(-120vh) scale(1.2); opacity: 0; }
    }
    </style>
    <script>
    // Add floating sparkles/particles
    if(!window.magicSparkleAdded){
        window.magicSparkleAdded = true;
        const sparkleLayer = document.createElement('div');
        sparkleLayer.className = 'magic-sparkle';
        for(let i=0;i<18;i++){
            const s = document.createElement('span');
            s.style.left = (Math.random()*100)+'vw';
            s.style.bottom = (-10-Math.random()*20)+'vh';
            s.style.width = (8+Math.random()*16)+'px';
            s.style.height = s.style.width;
            s.style.animationDelay = (Math.random()*6)+'s';
            sparkleLayer.appendChild(s);
        }
        document.body.appendChild(sparkleLayer);
    }
    </script>
''', unsafe_allow_html=True)

# --- ENHANCED MAGICAL BACKGROUND EFFECTS (NO AVATARS) ---
st.markdown('''
    <style>
    /* Animated floating lines (aurora waves) */
    .aurora-wave {
        position: fixed; left: 0; top: 0; width: 100vw; height: 100vh; z-index: 0;
        pointer-events: none;
        overflow: hidden;
    }
    .aurora-wave span {
        position: absolute;
        width: 120vw; height: 32px;
        left: -10vw;
        border-radius: 32px;
        opacity: 0.18;
        background: linear-gradient(90deg, #0057B8 0%, #FFD700 100%);
        filter: blur(12px);
        animation: auroramove 12s linear infinite;
    }
    .aurora-wave span:nth-child(2) {
        top: 30vh; animation-delay: 2s; opacity: 0.13; filter: blur(18px); }
    .aurora-wave span:nth-child(3) {
        top: 60vh; animation-delay: 4s; opacity: 0.11; filter: blur(22px); }
    .aurora-wave span:nth-child(4) {
        top: 80vh; animation-delay: 6s; opacity: 0.09; filter: blur(28px); }
    @keyframes auroramove {
        0% { left: -10vw; }
        100% { left: 10vw; }
    }
    /* Bokeh circles */
    .bokeh-bg {
        position: fixed; left: 0; top: 0; width: 100vw; height: 100vh; z-index: 0;
        pointer-events: none;
        overflow: hidden;
    }
    .bokeh-bg span {
        position: absolute;
        border-radius: 50%;
        opacity: 0.13;
        background: linear-gradient(135deg, #FFD700 0%, #0057B8 100%);
        filter: blur(8px);
        animation: bokehfloat 18s linear infinite;
    }
    @keyframes bokehfloat {
        0% { transform: translateY(0) scale(0.7); opacity: 0.13; }
        60% { opacity: 0.18; }
        100% { transform: translateY(-120vh) scale(1.2); opacity: 0; }
    }
    </style>
    <script>
    // Add aurora waves
    if(!window.auroraWaveAdded){
        window.auroraWaveAdded = true;
        const aurora = document.createElement('div');
        aurora.className = 'aurora-wave';
        for(let i=0;i<4;i++){
            const s = document.createElement('span');
            s.style.top = (10+20*i)+'vh';
            s.style.animationDuration = (10+4*i)+'s';
            aurora.appendChild(s);
        }
        document.body.appendChild(aurora);
    }
    // Add bokeh circles
    if(!window.bokehBgAdded){
        window.bokehBgAdded = true;
        const bokeh = document.createElement('div');
        bokeh.className = 'bokeh-bg';
        for(let i=0;i<12;i++){
            const s = document.createElement('span');
            s.style.left = (Math.random()*100)+'vw';
            s.style.bottom = (-10-Math.random()*20)+'vh';
            s.style.width = (32+Math.random()*48)+'px';
            s.style.height = s.style.width;
            s.style.animationDelay = (Math.random()*12)+'s';
            bokeh.appendChild(s);
        }
        document.body.appendChild(bokeh);
    }
    </script>
''', unsafe_allow_html=True)

# --- INTERACTIVE BACKGROUND: RIPPLE, WAVES, TUNABLE DENSITY/SPEED ---

# Add settings controls for density and speed
if "aurora_density" not in st.session_state:
    st.session_state["aurora_density"] = 4
if "bokeh_density" not in st.session_state:
    st.session_state["bokeh_density"] = 12
if "bg_speed" not in st.session_state:
    st.session_state["bg_speed"] = 1.0
if st.session_state["active_section"] == "Settings":
    st.markdown("**Background Effects**")
    st.session_state["aurora_density"] = st.slider("Aurora Wave Density", 1, 10, st.session_state["aurora_density"])
    st.session_state["bokeh_density"] = st.slider("Bokeh Circle Density", 2, 24, st.session_state["bokeh_density"])
    st.session_state["bg_speed"] = st.slider("Background Animation Speed", 0.5, 2.0, st.session_state["bg_speed"], step=0.05)

# Interactive background JS/CSS
st.markdown(f'''
    <style>
    .ripple-bg {{ position: fixed; left: 0; top: 0; width: 100vw; height: 100vh; z-index: 2; pointer-events: none; }}
    .ripple-bg span {{ position: absolute; border-radius: 50%; opacity: 0.22; background: radial-gradient(circle, #0057B8 0%, #FFD700 80%, transparent 100%); pointer-events: none; animation: ripplescale 1.2s cubic-bezier(.4,2,.6,1); }}
    @keyframes ripplescale {{
        0% {{ transform: scale(0.2); opacity: 0.5; }}
        60% {{ opacity: 0.22; }}
        100% {{ transform: scale(2.2); opacity: 0; }}
    }}
    .cursor-wave {{ position: fixed; left: 0; top: 0; width: 100vw; height: 100vh; z-index: 2; pointer-events: none; }}
    .cursor-wave span {{ position: absolute; width: 48px; height: 48px; border-radius: 50%; opacity: 0.13; background: linear-gradient(135deg, #FFD700 0%, #0057B8 100%); filter: blur(8px); pointer-events: none; animation: wavefade 1.2s linear; }}
    @keyframes wavefade {{ 0% {{ opacity: 0.13; }} 100% {{ opacity: 0; }} }}
    </style>
    <script>
    // Remove old aurora/bokeh if present
    if(window.auroraWaveAdded){{
        document.querySelectorAll('.aurora-wave').forEach(e=>e.remove());
        window.auroraWaveAdded = false;
    }}
    if(window.bokehBgAdded){{
        document.querySelectorAll('.bokeh-bg').forEach(e=>e.remove());
        window.bokehBgAdded = false;
    }}
    // Add aurora waves
    if(!window.auroraWaveAdded){{
        window.auroraWaveAdded = true;
        const aurora = document.createElement('div');
        aurora.className = 'aurora-wave';
        const density = {st.session_state.get('aurora_density', 4)};
        const speed = {st.session_state.get('bg_speed', 1.0)};
        for(let i=0;i<density;i++){{
            const s = document.createElement('span');
            s.style.top = (10+80*i/density)+'vh';
            s.style.animationDuration = (10+4*i)*(1/speed)+'s';
            aurora.appendChild(s);
        }}
        document.body.appendChild(aurora);
    }}
    // Add bokeh circles
    if(!window.bokehBgAdded){{
        window.bokehBgAdded = true;
        const bokeh = document.createElement('div');
        bokeh.className = 'bokeh-bg';
        const bokehDensity = {st.session_state.get('bokeh_density', 12)};
        const speed = {st.session_state.get('bg_speed', 1.0)};
        for(let i=0;i<bokehDensity;i++){{
            const s = document.createElement('span');
            s.style.left = (Math.random()*100)+'vw';
            s.style.bottom = (-10-Math.random()*20)+'vh';
            s.style.width = (32+Math.random()*48)+'px';
            s.style.height = s.style.width;
            s.style.animationDelay = (Math.random()*12)+'s';
            s.style.animationDuration = (18*(1/speed))+'s';
            bokeh.appendChild(s);
        }}
        document.body.appendChild(bokeh);
    }}
    // Add ripple effect on click
    if(!window.rippleBgAdded){{
        window.rippleBgAdded = true;
        const ripple = document.createElement('div');
        ripple.className = 'ripple-bg';
        document.body.appendChild(ripple);
        window.parent.document.body.addEventListener('click', function(e){{
            if(e.target.closest('.stApp')){{
                const span = document.createElement('span');
                span.style.left = (e.clientX-80)+'px';
                span.style.top = (e.clientY-80)+'px';
                span.style.width = '160px';
                span.style.height = '160px';
                ripple.appendChild(span);
                setTimeout(()=>{{span.remove();}}, 1200);
            }}
        }});
    }}
    // Add cursor-following wave
    if(!window.cursorWaveAdded){{
        window.cursorWaveAdded = true;
        const wave = document.createElement('div');
        wave.className = 'cursor-wave';
        document.body.appendChild(wave);
        window.parent.document.body.addEventListener('mousemove', function(e){{
            if(e.target.closest('.stApp')){{
                const span = document.createElement('span');
                span.style.left = (e.clientX-24)+'px';
                span.style.top = (e.clientY-24)+'px';
                wave.appendChild(span);
                setTimeout(()=>{{span.remove();}}, 1200);
            }}
        }});
    }}
    </script>
''', unsafe_allow_html=True)
